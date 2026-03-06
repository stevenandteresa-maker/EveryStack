#!/usr/bin/env python3
"""
Generate the Phase 2B Operator Guide (.docx)
Color-coded, container-formatted runbook for Steven.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, color_hex="999999"):
    """Set borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'</w:tcBorders>'
    )
    tcBorders_existing = tcPr.find(qn('w:tcBorders'))
    if tcBorders_existing is not None:
        tcPr.remove(tcBorders_existing)
    tcPr.append(tcBorders)


def add_container(doc, text, bg_color, border_color=None, is_code=False, bold_first_line=False):
    """Add a colored container (1-cell table) with text."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg_color)
    if border_color:
        set_cell_border(cell, border_color)
    else:
        set_cell_border(cell, bg_color)

    lines = text.strip().split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()

        p.space_before = Pt(2)
        p.space_after = Pt(2)

        run = p.add_run(line)
        if is_code:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

        if bold_first_line and i == 0:
            run.bold = True

    # Add spacing after
    doc.add_paragraph()


def add_gray_box(doc, title, what_builds, what_youll_see, how_long):
    """Gray explanation container."""
    text = f"{title}\n\nWhat This Builds:\n{what_builds}\n\nWhat You'll See When It's Done:\n{what_youll_see}\n\nHow Long This Typically Takes:\n{how_long}"
    add_container(doc, text, "E8E8E8", "CCCCCC", bold_first_line=True)


def add_blue_box(doc, text):
    """Blue prompt container — paste into Claude Code."""
    add_container(doc, text, "DBEAFE", "93C5FD", is_code=True)


def add_green_box(doc, text):
    """Green git action container."""
    add_container(doc, text, "D1FAE5", "6EE7B7", is_code=True)


def add_orange_box(doc, text):
    """Orange checkpoint container."""
    add_container(doc, text, "FED7AA", "FDBA74", bold_first_line=True)


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def main():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ══════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('Phase 2B — Operator Guide', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    subtitle = doc.add_heading('Synced Data Performance, Outbound Sync,\nConflict Resolution', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('EveryStack Build Runbook')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()
    doc.add_paragraph()

    # Color key
    add_heading(doc, 'Color Key', level=2)
    p = doc.add_paragraph()
    p.add_run('This guide uses four colors:').font.size = Pt(10)

    add_container(doc, 'BLUE — Prompt to paste into Claude Code', 'DBEAFE', '93C5FD')
    add_container(doc, 'GREEN — Git command to run in your terminal', 'D1FAE5', '6EE7B7')
    add_container(doc, 'ORANGE — Checkpoint: verify before moving on', 'FED7AA', 'FDBA74')
    add_container(doc, 'GRAY — Explanation: read this, don\'t paste it', 'E8E8E8', 'CCCCCC')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # SETUP
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'SETUP', level=1)

    add_gray_box(doc,
        'Before You Begin',
        'Nothing yet — this is setup. You are getting your environment ready for Phase 2B.',
        'Your terminal will be open in the EveryStack monorepo, on a new feature branch, with Claude Code running.',
        '2–3 minutes.'
    )

    p = doc.add_paragraph('Follow these steps to set up your environment:')
    p.space_after = Pt(4)

    steps = [
        '1. Open your terminal (or VS Code integrated terminal).',
        '2. Navigate to the EveryStack monorepo:',
    ]
    for s in steps:
        doc.add_paragraph(s)

    add_blue_box(doc, 'cd ~/Documents/EveryStack')

    doc.add_paragraph('3. Make sure you are on the main branch and up to date:')
    add_green_box(doc, 'git checkout main && git pull origin main')

    doc.add_paragraph('4. Create the Phase 2B feature branch:')
    add_green_box(doc, 'git checkout -b feat/phase-2b-outbound-sync-conflicts')

    doc.add_paragraph('5. Open Claude Code (in VS Code or your terminal).')
    doc.add_paragraph('6. Load the required skill files. Paste this into Claude Code:')

    add_blue_box(doc,
        'Read these skill files and keep their conventions in mind for all work in this session:\n'
        '- docs/skills/backend/SKILL.md\n'
        '- docs/skills/ux-ui/SKILL.md\n'
        '- docs/skills/phase-context/SKILL.md'
    )

    doc.add_paragraph('You are now ready to begin. Follow each prompt below in order.')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # PROMPT 1
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'PROMPT 1: Speed Up Grid Queries on Synced Data', level=1)

    add_gray_box(doc,
        'PROMPT 1 — What This Step Does',
        'This builds database indexes that make sorting and filtering fast on synced tables. Think of it like adding an index to the back of a textbook — instead of reading every page to find something, the database can jump straight to the right data. This is what allows grid views to load in under 200 milliseconds even with 50,000 records.',
        'Claude Code will create two new files: a database migration file and a utility file for managing indexes. It will also add query helper functions. You should see tests pass with green checkmarks. If any test fails, Claude Code will attempt to fix it automatically.',
        '5–8 minutes.'
    )

    add_blue_box(doc,
        'Read docs/reference/database-scaling.md lines 486–498 and docs/reference/sync-engine.md lines 437–443 for context.\n'
        '\n'
        'Build the JSONB expression index infrastructure for grid query performance on synced tables.\n'
        '\n'
        'Target files:\n'
        '- packages/shared/db/migrations/XXXX_jsonb_expression_indexes.ts\n'
        '- packages/shared/db/index-utils.ts\n'
        '- packages/shared/db/query-helpers.ts (extend if exists, create if not)\n'
        '\n'
        'What to build:\n'
        '\n'
        '1. A utility function `createFieldExpressionIndex(tenantId, tableId, fieldId, fieldType)` that generates and executes CREATE INDEX CONCURRENTLY statements on canonical_data JSONB paths. Index names must be deterministic: `idx_rec_{short_hash(tenantId, tableId, fieldId)}` and must not exceed PostgreSQL\'s 63-character limit.\n'
        '\n'
        '2. A complementary `dropFieldExpressionIndex(tenantId, tableId, fieldId)` function.\n'
        '\n'
        '3. Query helpers: `canonicalFieldExpression(fieldId, fieldType)` that returns Drizzle-compatible SQL for WHERE clauses, and `canonicalFieldOrderBy(fieldId, fieldType, direction)` for ORDER BY clauses. Both must use proper type casting (numeric for numbers, text for text).\n'
        '\n'
        '4. A migration file that uses CREATE INDEX CONCURRENTLY (NOT inside a transaction — CONCURRENTLY requires no wrapping transaction). Raw SQL is acceptable here since Drizzle does not support CREATE INDEX CONCURRENTLY or expression indexes.\n'
        '\n'
        '5. Add a code comment in index-utils.ts referencing the decision point: "If expression indexes cannot meet <200ms grid query targets on 50K+ record tables, introduce the record_cells denormalized read cache (see data-model.md)."\n'
        '\n'
        'Support these field types in the expression/order helpers: text, number, date, single_select, checkbox.\n'
        '\n'
        'Acceptance criteria:\n'
        '- createFieldExpressionIndex() generates valid CREATE INDEX CONCURRENTLY\n'
        '- dropFieldExpressionIndex() drops the matching index\n'
        '- Index names are deterministic and under 63 chars\n'
        '- canonicalFieldExpression() returns correct SQL for all 5 field types\n'
        '- canonicalFieldOrderBy() returns correct SQL with proper type casting\n'
        '- Unit tests cover all 5 field type expressions\n'
        '- Integration test creates an index and verifies EXPLAIN shows index scan\n'
        '- testTenantIsolation() passes\n'
        '- ESLint and TypeScript compile with zero errors\n'
        '- Coverage >= 80% on new files\n'
        '\n'
        'Do NOT build: automatic index suggestion, record_cells denormalization, grid view queries, per-tenant index management UI.'
    )

    add_orange_box(doc,
        'Checkpoint — Verify Prompt 1\n'
        'Look for:\n'
        '- Two new files created (migration + index-utils.ts)\n'
        '- Query helpers in query-helpers.ts\n'
        '- All tests passing (green checkmarks)\n'
        '- No TypeScript or ESLint errors'
    )

    add_green_box(doc,
        'git add packages/shared/db/migrations/ packages/shared/db/index-utils.ts packages/shared/db/query-helpers.ts\n'
        'git commit -m "feat(db): add JSONB expression indexes on canonical_data for grid query performance [Phase 2B, Prompt 1]"'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # PROMPT 2
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'PROMPT 2: Track Last-Synced Values for Conflict Detection', level=1)

    add_gray_box(doc,
        'PROMPT 2 — What This Step Does',
        'This defines how EveryStack remembers what each field\'s value was the last time it successfully synced with the external platform (Airtable, etc.). This "last synced value" is the reference point needed to detect conflicts later — if both EveryStack and Airtable changed a field since the last sync, that\'s a conflict.',
        'Claude Code will add new types, utility functions, and update the inbound sync pipeline. You should see new files in the sync package and updated sync worker code. All tests should pass.',
        '5–8 minutes.'
    )

    add_blue_box(doc,
        'Read docs/reference/sync-engine.md lines 543–549 and lines 30–65, and docs/reference/data-model.md lines 104–112 for context.\n'
        '\n'
        'Define the sync_metadata JSONB shape and integrate last_synced_value tracking into the inbound sync pipeline.\n'
        '\n'
        'Target files:\n'
        '- packages/shared/sync/types.ts (extend)\n'
        '- packages/shared/sync/sync-metadata.ts (new)\n'
        '- apps/worker/src/jobs/sync-inbound.ts (extend)\n'
        '\n'
        'The records table already has a sync_metadata JSONB column from Phase 1B. This prompt defines its shape and populates it.\n'
        '\n'
        'What to build:\n'
        '\n'
        '1. Define the SyncMetadata interface in types.ts:\n'
        '   - last_synced_at: ISO 8601 timestamp\n'
        '   - last_synced_values: Record<string, { value: CanonicalFieldValue, synced_at: string }>\n'
        '   - sync_status: \'active\' | \'orphaned\'\n'
        '   - sync_direction: \'inbound\' | \'outbound\' | \'both\'\n'
        '\n'
        '2. Build sync-metadata.ts with utility functions:\n'
        '   - updateLastSyncedValues(existingMetadata, updatedFieldIds, canonicalData): SyncMetadata\n'
        '   - getLastSyncedValue(metadata, fieldId): CanonicalFieldValue | undefined\n'
        '   - createInitialSyncMetadata(canonicalData, fieldIds): SyncMetadata\n'
        '\n'
        '3. Update the inbound sync pipeline (sync-inbound.ts):\n'
        '   - For new records: call createInitialSyncMetadata() and write to the record\n'
        '   - For updated records: call updateLastSyncedValues() for changed fields\n'
        '   - sync_metadata update and canonical_data update must be in the SAME database transaction\n'
        '\n'
        '4. Add a Zod schema for SyncMetadata to validate the JSONB shape on read.\n'
        '\n'
        'Acceptance criteria:\n'
        '- SyncMetadata type exported from types.ts\n'
        '- createInitialSyncMetadata() populates all synced field values with timestamps\n'
        '- updateLastSyncedValues() updates only specified fields, preserves others\n'
        '- getLastSyncedValue() returns correct value or undefined\n'
        '- Inbound sync writes sync_metadata for new and updated records\n'
        '- sync_metadata and canonical_data updates in same transaction\n'
        '- Zod schema validates the shape\n'
        '- testTenantIsolation() passes\n'
        '- Unit tests with edge cases (empty metadata, partial updates, missing fields)\n'
        '- ESLint and TypeScript compile with zero errors\n'
        '- Coverage >= 80% on new files\n'
        '\n'
        'Do NOT build: conflict detection logic, outbound sync pipeline, sync_metadata for native tables, field-level sync history.'
    )

    add_orange_box(doc,
        'Checkpoint — Verify Prompt 2\n'
        'Look for:\n'
        '- New file: packages/shared/sync/sync-metadata.ts\n'
        '- Updated types in packages/shared/sync/types.ts\n'
        '- Updated sync-inbound.ts with sync_metadata writes\n'
        '- All tests passing'
    )

    add_green_box(doc,
        'git add packages/shared/sync/types.ts packages/shared/sync/sync-metadata.ts apps/worker/src/jobs/sync-inbound.ts\n'
        'git commit -m "feat(sync): define sync_metadata JSONB shape and last_synced_value tracking [Phase 2B, Prompt 2]"'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # PROMPT 3
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'PROMPT 3: Build the Outbound Sync Pipeline', level=1)

    add_gray_box(doc,
        'PROMPT 3 — What This Step Does',
        'This builds the pipeline that pushes EveryStack edits back to the source platform (e.g., Airtable). When a user edits a cell in a synced table, the change is saved locally first (fast), then a background job sends the update to Airtable. This prompt builds that background job and the rate-limiting logic to avoid hitting API limits.',
        'Claude Code will create the outbound sync worker job, the core sync logic, and the enqueue helper. You should see new files in the sync package and worker. Tests will run against mock API responses.',
        '8–12 minutes.'
    )

    add_blue_box(doc,
        'Read docs/reference/sync-engine.md lines 446–453 and lines 30–43, and docs/reference/data-model.md lines 649–660 for context.\n'
        '\n'
        'Build the outbound sync pipeline that pushes EveryStack edits to the source platform via BullMQ.\n'
        '\n'
        'Target files:\n'
        '- apps/worker/src/jobs/sync-outbound.ts (new)\n'
        '- packages/shared/sync/outbound.ts (new)\n'
        '- packages/shared/sync/types.ts (extend with OutboundSyncJob, OutboundSyncResult types)\n'
        '\n'
        'What to build:\n'
        '\n'
        '1. Define OutboundSyncJob and OutboundSyncResult types in types.ts with fields: tenantId, recordId, tableId, baseConnectionId, changedFieldIds, editedBy, priority.\n'
        '\n'
        '2. Build executeOutboundSync(job) in outbound.ts that:\n'
        '   - Reads the record\'s canonical_data and source_refs\n'
        '   - Looks up synced_field_mappings for external_field_id and external_field_type\n'
        '   - Calls fromCanonical() via FieldTypeRegistry to transform back to platform format\n'
        '   - Acquires a rate limit token from the Redis token-bucket (built in Phase 2A)\n'
        '   - Calls the platform API to update the record\n'
        '   - On success: updates sync_metadata.last_synced_values\n'
        '   - On failure: returns error details (does NOT retry — BullMQ handles retries)\n'
        '\n'
        '3. Build the BullMQ job processor in sync-outbound.ts:\n'
        '   - Queue: sync:outbound\n'
        '   - Retry: 3 attempts with exponential backoff (1min, 5min, 15min)\n'
        '   - On final failure: log via Pino + Sentry\n'
        '\n'
        '4. Build enqueueOutboundSync(tenantId, recordId, tableId, changedFieldIds, editedBy, priority?) that:\n'
        '   - Checks if the table is synced (looks up base_connection)\n'
        '   - If not synced: returns immediately (no-op)\n'
        '   - If synced: adds BullMQ job\n'
        '   - Deduplicates: merges changedFieldIds if a job for the same recordId is already queued\n'
        '\n'
        '5. Skip computed fields: never sync Lookup, Rollup, Formula, Count back to platforms (check isLossless in the registry).\n'
        '\n'
        'Acceptance criteria:\n'
        '- executeOutboundSync() transforms via fromCanonical() and writes to platform API\n'
        '- Rate limit token acquired before API call\n'
        '- sync_metadata.last_synced_values updated on success\n'
        '- Computed fields excluded from outbound\n'
        '- BullMQ retry with exponential backoff (3 attempts)\n'
        '- enqueueOutboundSync() deduplicates and is no-op for native tables\n'
        '- Unit tests: successful sync, rate limit wait, API error, computed field skip, dedup\n'
        '- Integration test with MSW mock for end-to-end outbound flow\n'
        '- testTenantIsolation() passes\n'
        '- ESLint and TypeScript compile with zero errors\n'
        '- Coverage >= 85% on outbound.ts\n'
        '\n'
        'Do NOT build: optimistic UI (Prompt 4), conflict detection on failures (Prompt 5), P0-P3 scheduling (Phase 2C), Notion/SmartSuite outbound (only Airtable).'
    )

    add_orange_box(doc,
        'Checkpoint — Verify Prompt 3\n'
        'Look for:\n'
        '- New files: sync-outbound.ts (worker) and outbound.ts (shared)\n'
        '- Extended types in types.ts\n'
        '- Tests passing with mock API responses\n'
        '- No TypeScript or ESLint errors'
    )

    add_green_box(doc,
        'git add apps/worker/src/jobs/sync-outbound.ts packages/shared/sync/outbound.ts packages/shared/sync/types.ts\n'
        'git commit -m "feat(sync): outbound sync BullMQ job with fromCanonical() and rate-limited platform write [Phase 2B, Prompt 3]"'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # PROMPT 4
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'PROMPT 4: Make Synced Table Cells Editable', level=1)

    add_gray_box(doc,
        'PROMPT 4 — What This Step Does',
        'This builds the server-side action that lets users edit a cell in a synced table. The edit saves instantly to EveryStack\'s database (so it feels fast), then queues a background job to push the change to the external platform. It also protects computed fields (like formulas and lookups) from being edited, since those are read-only from the source platform.',
        'Claude Code will create a Server Action file and extend the records data layer. Tests will verify that edits work, computed fields are protected, and tenant isolation is enforced.',
        '5–8 minutes.'
    )

    add_blue_box(doc,
        'Read docs/reference/sync-engine.md lines 446–453 and docs/reference/data-model.md lines 649–660 for context.\n'
        '\n'
        'Build the Server Action for editing a cell in a synced table with optimistic local update and background outbound sync.\n'
        '\n'
        'Target files:\n'
        '- apps/web/src/actions/sync-edit.ts (new)\n'
        '- apps/web/src/data/records.ts (extend)\n'
        '\n'
        'What to build:\n'
        '\n'
        '1. Server Action updateSyncedRecordField that:\n'
        '   - Validates input via Zod: { recordId, fieldId, newValue, tableId }\n'
        '   - Verifies the field is editable (not computed, not read-only per isLossless: false)\n'
        '   - Updates canonical_data in PostgreSQL immediately via getDbForTenant(tenantId, \'write\')\n'
        '   - Updates search_vector via buildSearchVector() in the SAME transaction\n'
        '   - Calls enqueueOutboundSync() to queue the platform write\n'
        '   - Returns the updated record immediately\n'
        '\n'
        '2. Helper getOutboundSyncStatus(recordId, fieldId) in records.ts that returns \'synced\' | \'pending\' | \'failed\' by checking BullMQ job state.\n'
        '\n'
        '3. If user tries to edit a computed field (Lookup, Rollup, Formula, Count): return validation error { code: \'VALIDATION_FAILED\', message: \'This field is synced from {platform} and cannot be edited.\' }\n'
        '\n'
        'Acceptance criteria:\n'
        '- updateSyncedRecordField updates canonical_data and returns immediately\n'
        '- search_vector updated in same transaction\n'
        '- enqueueOutboundSync() called after local update\n'
        '- Computed fields return validation error\n'
        '- getOutboundSyncStatus() reports synced/pending/failed correctly\n'
        '- Zod validation rejects malformed input\n'
        '- testTenantIsolation() passes\n'
        '- Unit tests: successful edit, computed field rejection, invalid input\n'
        '- Integration test: local update + outbound job enqueued\n'
        '- ESLint and TypeScript compile with zero errors\n'
        '- Coverage >= 90% on sync-edit.ts\n'
        '\n'
        'Do NOT build: cell renderer UI or inline editing (Phase 3), client-side optimistic state (Phase 3), conflict detection on outbound failure (Prompt 5), real-time push of edit events (Phase 3).'
    )

    add_orange_box(doc,
        'Checkpoint — Verify Prompt 4\n'
        'Look for:\n'
        '- New file: apps/web/src/actions/sync-edit.ts\n'
        '- Extended records.ts with getOutboundSyncStatus()\n'
        '- All tests passing\n'
        '- No TypeScript or ESLint errors'
    )

    add_green_box(doc,
        'git add apps/web/src/actions/sync-edit.ts apps/web/src/data/records.ts\n'
        'git commit -m "feat(sync): optimistic UI for synced table cell edits with outbound sync enqueue [Phase 2B, Prompt 4]"'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # INTEGRATION CHECKPOINT 1
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'INTEGRATION CHECKPOINT 1', level=1)

    add_gray_box(doc,
        'Integration Checkpoint 1 — Verify Prompts 1–4',
        'This is a verification step. You are running the full test suite and checks to make sure everything from Prompts 1 through 4 works together correctly before moving on to the conflict detection section.',
        'You will run several commands and see test results. Everything should pass with green checkmarks. If anything fails, paste the error into Claude Code and let it fix the issue before continuing.',
        '3–5 minutes.'
    )

    add_blue_box(doc,
        'Run the full verification suite for Prompts 1–4:\n'
        '\n'
        'pnpm turbo typecheck\n'
        'pnpm turbo lint\n'
        'pnpm turbo test\n'
        'pnpm turbo test -- --coverage\n'
        'pnpm turbo db:migrate:check\n'
        '\n'
        'Then verify:\n'
        '- createFieldExpressionIndex() can create an index on a test field and EXPLAIN shows index scan\n'
        '- Inbound sync writes sync_metadata.last_synced_values to records\n'
        '- updateSyncedRecordField() updates canonical_data and enqueues outbound sync job\n'
        '- Outbound sync job processes and calls the Airtable API (via MSW mock)\n'
        '\n'
        'Fix any failures before proceeding.'
    )

    add_orange_box(doc,
        'Checkpoint — All Must Pass\n'
        '- TypeScript: zero errors\n'
        '- ESLint: zero errors\n'
        '- Tests: all green\n'
        '- Coverage: thresholds met\n'
        '- Migration check: no lock violations\n'
        '\n'
        'If anything fails, paste the error into Claude Code and say "fix this". Do NOT proceed until everything passes.'
    )

    add_green_box(doc,
        'git add -A\n'
        'git commit -m "chore(verify): integration checkpoint 1 — JSONB indexes, sync_metadata, outbound sync [Phase 2B, CP-1]"\n'
        'git push -u origin feat/phase-2b-outbound-sync-conflicts'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # PROMPT 5
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'PROMPT 5: Build Conflict Detection', level=1)

    add_gray_box(doc,
        'PROMPT 5 — What This Step Does',
        'This builds the algorithm that detects when both EveryStack and the external platform have changed the same field since the last sync. It uses a "three-way comparison": (1) what EveryStack has now, (2) what the platform is sending, and (3) what the value was at last sync. If both sides diverged from the last-synced value, that\'s a conflict. This is the brain of the conflict system.',
        'Claude Code will create the conflict detection module and integrate it into the inbound sync pipeline. You should see a new conflict-detection.ts file and updated sync-inbound.ts. Tests will cover all possible comparison outcomes.',
        '8–12 minutes.'
    )

    add_blue_box(doc,
        'Read docs/reference/sync-engine.md lines 535–570 for context on conflict detection.\n'
        '\n'
        'Build the three-way conflict detection algorithm and integrate it into the inbound sync pipeline.\n'
        '\n'
        'Target files:\n'
        '- packages/shared/sync/conflict-detection.ts (new)\n'
        '- apps/worker/src/jobs/sync-inbound.ts (extend)\n'
        '- packages/shared/sync/types.ts (extend)\n'
        '\n'
        'What to build:\n'
        '\n'
        '1. Build conflict-detection.ts with:\n'
        '   - ConflictDetectionResult, DetectedConflict, CleanChange interfaces\n'
        '   - detectConflicts(currentCanonical, inboundCanonical, syncMetadata, syncedFieldIds) function\n'
        '\n'
        '   Three-way comparison per field:\n'
        '   - baseValue = syncMetadata.last_synced_values[fieldId].value\n'
        '   - localValue = currentCanonical.fields[fieldId]\n'
        '   - remoteValue = inboundCanonical.fields[fieldId]\n'
        '   Results:\n'
        '   - local unchanged + remote changed → clean remote change (apply it)\n'
        '   - local changed + remote unchanged → clean local change (keep it)\n'
        '   - both changed to DIFFERENT values → CONFLICT\n'
        '   - both changed to SAME value → convergent (no conflict)\n'
        '   - all three equal → unchanged\n'
        '\n'
        '2. Integrate into inbound sync pipeline (sync-inbound.ts):\n'
        '   - For each record: call detectConflicts()\n'
        '   - Clean remote changes: apply to canonical_data, update sync_metadata\n'
        '   - Conflicts: write sync_conflicts records (status: pending)\n'
        '   - Clean local changes: no action (preserved)\n'
        '\n'
        '3. Build writeConflictRecords(tx, tenantId, recordId, conflicts, platform) that writes one sync_conflicts row per conflicted field.\n'
        '\n'
        '4. Edge cases:\n'
        '   - If sync_metadata is null or last_synced_values empty: treat as clean remote change\n'
        '   - New records from platform: create normally, no conflict possible\n'
        '   - Null and undefined are equivalent in comparisons\n'
        '\n'
        'Acceptance criteria:\n'
        '- detectConflicts() identifies conflicts, clean changes, and unchanged fields correctly\n'
        '- Convergent changes are NOT flagged as conflicts\n'
        '- sync_conflicts records written with correct local_value, remote_value, base_value\n'
        '- Clean remote changes applied to canonical_data\n'
        '- Clean local changes preserved\n'
        '- Records without sync_metadata treat inbound as clean changes\n'
        '- testTenantIsolation() passes\n'
        '- Unit tests: all 5 comparison outcomes, null handling, convergent, missing metadata\n'
        '- Integration test: inbound sync with mixed conflicts and clean changes\n'
        '- ESLint and TypeScript compile with zero errors\n'
        '- Coverage >= 90% on conflict-detection.ts\n'
        '\n'
        'Do NOT build: conflict resolution UI (Prompts 6-7, 10), grid rendering (Prompt 8), real-time push (Prompt 9), auto-resolution (Prompt 6).'
    )

    add_orange_box(doc,
        'Checkpoint — Verify Prompt 5\n'
        'Look for:\n'
        '- New file: packages/shared/sync/conflict-detection.ts\n'
        '- Updated sync-inbound.ts with conflict detection integration\n'
        '- Tests covering all 5 comparison outcomes\n'
        '- All tests passing'
    )

    add_green_box(doc,
        'git add packages/shared/sync/conflict-detection.ts packages/shared/sync/types.ts apps/worker/src/jobs/sync-inbound.ts\n'
        'git commit -m "feat(sync): three-way conflict detection on inbound sync with sync_conflicts population [Phase 2B, Prompt 5]"'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # PROMPT 6
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'PROMPT 6: Set Up Conflict Resolution Strategy', level=1)

    add_gray_box(doc,
        'PROMPT 6 — What This Step Does',
        'This builds the default way EveryStack handles conflicts: "last write wins" — meaning the external platform\'s value is automatically applied, and the old local value is preserved for recovery. It also builds a toggle that lets Managers switch a synced table to "manual mode", where conflicts stay unresolved until a human decides which value to keep.',
        'Claude Code will create the conflict resolution logic and a toggle Server Action. Tests will verify both auto-resolution and manual mode.',
        '5–8 minutes.'
    )

    add_blue_box(doc,
        'Read docs/reference/sync-engine.md lines 571–608 for context.\n'
        '\n'
        'Build the default conflict resolution strategy and manual resolution mode toggle.\n'
        '\n'
        'Target files:\n'
        '- packages/shared/sync/conflict-resolution.ts (new)\n'
        '- apps/web/src/actions/sync-settings.ts (new)\n'
        '- apps/web/src/data/sync-settings.ts (new)\n'
        '\n'
        'What to build:\n'
        '\n'
        '1. Build conflict-resolution.ts with applyLastWriteWins(tx, tenantId, recordId, conflicts, platform):\n'
        '   - For each conflict: apply remote value to canonical_data (inbound wins)\n'
        '   - Create sync_conflicts records with status: \'resolved_remote\' (preserving local value)\n'
        '   - Update sync_metadata.last_synced_values with remote values\n'
        '   - Log: "Auto-resolved {count} conflicts via last-write-wins for record {recordId}"\n'
        '\n'
        '2. When manual_conflict_resolution is true:\n'
        '   - Conflicts written with status: \'pending\'\n'
        '   - Inbound value NOT applied to canonical_data (local preserved)\n'
        '   - sync_metadata.last_synced_values NOT updated for conflicted fields\n'
        '\n'
        '3. Update the inbound sync pipeline:\n'
        '   - After detectConflicts(), check base_connections.sync_config.manual_conflict_resolution\n'
        '   - If false/missing: call applyLastWriteWins()\n'
        '   - If true: write sync_conflicts with pending status, preserve local data\n'
        '\n'
        '4. Server Action toggleManualConflictResolution in sync-settings.ts:\n'
        '   - Input: { baseConnectionId, tableId, enabled: boolean }\n'
        '   - Permission: Owner, Admin, or Manager\n'
        '   - Updates base_connections.sync_config JSONB\n'
        '\n'
        '5. Data function getConflictResolutionMode in sync-settings.ts.\n'
        '\n'
        'Acceptance criteria:\n'
        '- Last-write-wins: remote value applied, sync_conflicts status resolved_remote\n'
        '- Overwritten local values preserved in sync_conflicts.local_value\n'
        '- Manual mode: status pending, local canonical_data unchanged\n'
        '- Toggle requires Manager+ permission\n'
        '- testTenantIsolation() passes\n'
        '- Unit tests: LWW flow, manual mode, toggle with permission check\n'
        '- ESLint and TypeScript compile with zero errors\n'
        '- Coverage >= 90% on conflict-resolution.ts\n'
        '\n'
        'Do NOT build: resolution modal UI (Prompt 7), grid rendering (Prompt 8), Keep EveryStack/Edit actions (Prompt 10), bulk resolution (Prompt 11).'
    )

    add_orange_box(doc,
        'Checkpoint — Verify Prompt 6\n'
        'Look for:\n'
        '- New files: conflict-resolution.ts, sync-settings.ts (action + data)\n'
        '- Updated inbound sync pipeline with mode check\n'
        '- Tests for both auto and manual modes\n'
        '- All tests passing'
    )

    add_green_box(doc,
        'git add packages/shared/sync/conflict-resolution.ts apps/web/src/actions/sync-settings.ts apps/web/src/data/sync-settings.ts apps/worker/src/jobs/sync-inbound.ts\n'
        'git commit -m "feat(sync): default last-write-wins resolution and manual resolution mode toggle [Phase 2B, Prompt 6]"'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # PROMPT 7
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'PROMPT 7: Build the Conflict Resolution Popup', level=1)

    add_gray_box(doc,
        'PROMPT 7 — What This Step Does',
        'This builds the popup (modal) that Managers see when they click on a conflicted cell to resolve it. For a single conflict, it shows both values side by side with buttons to choose. For multiple conflicts on the same record, it shows a scrollable list with per-field choices and "Keep All" bulk buttons. All text will be translatable (i18n).',
        'Claude Code will create three React component files. You should see the modal, field row, and action bar components. Component tests will verify rendering with mock data.',
        '8–12 minutes.'
    )

    add_blue_box(doc,
        'Read docs/reference/sync-engine.md lines 575–681 for the modal wireframes and resolution actions.\n'
        '\n'
        'Build the conflict resolution modal UI.\n'
        '\n'
        'Target files:\n'
        '- apps/web/src/components/sync/ConflictResolutionModal.tsx (new)\n'
        '- apps/web/src/components/sync/ConflictFieldRow.tsx (new)\n'
        '- apps/web/src/components/sync/ConflictResolutionActions.tsx (new)\n'
        '\n'
        'What to build:\n'
        '\n'
        '1. ConflictResolutionModal.tsx — shadcn/ui Dialog-based modal:\n'
        '   - Single-field mode: shows field name, EveryStack value vs Platform value side by side, base value ("Was: ..."), and 3 buttons: Keep EveryStack, Keep {Platform}, Edit\n'
        '   - Multi-field mode: scrollable list of ConflictFieldRow components + bulk action bar\n'
        '   - Once resolved, row shows green checkmark and chosen value\n'
        '\n'
        '2. ConflictFieldRow.tsx — reusable row for one field\'s conflict:\n'
        '   - Shows: field name, local value, remote value, base value, who changed/when\n'
        '   - 3 action buttons per field\n'
        '   - "Edit" opens inline text/value input for merged value (basic for now — Phase 3 field editors replace this)\n'
        '\n'
        '3. ConflictResolutionActions.tsx — bulk action bar:\n'
        '   - "Keep All EveryStack" and "Keep All {Platform}" buttons\n'
        '   - Disabled when all conflicts on the record are already resolved\n'
        '\n'
        '4. All user-facing strings via i18n: useTranslations(\'SyncConflicts\'). Add keys to messages/en.json.\n'
        '\n'
        '5. Render field values using FieldTypeRegistry display formatters (not raw JSON).\n'
        '\n'
        'Acceptance criteria:\n'
        '- Single-field modal renders with local, remote, base values and action buttons\n'
        '- Multi-field modal renders scrollable list + bulk action bar\n'
        '- Keep EveryStack, Keep {Platform}, Edit buttons present on each row\n'
        '- Bulk Keep All buttons resolve all pending conflicts\n'
        '- Field values rendered via FieldTypeRegistry (not raw JSON)\n'
        '- All text uses i18n\n'
        '- shadcn/ui Dialog, DM Sans font, 4px spacing multiples\n'
        '- Touch targets >= 44x44px (WCAG 2.5.8)\n'
        '- Component tests verify rendering with mock data\n'
        '- ESLint and TypeScript compile with zero errors\n'
        '- Coverage >= 80% on new files\n'
        '\n'
        'Do NOT build: the Server Action that saves resolution decisions (Prompt 10), grid indicators (Prompt 8), real-time push (Prompt 9), mobile layout (Phase 3H), optimistic resolution with undo (Prompt 10).'
    )

    add_orange_box(doc,
        'Checkpoint — Verify Prompt 7\n'
        'Look for:\n'
        '- Three new component files in apps/web/src/components/sync/\n'
        '- i18n keys added to messages/en.json\n'
        '- Component tests passing\n'
        '- No TypeScript or ESLint errors'
    )

    add_green_box(doc,
        'git add apps/web/src/components/sync/ messages/en.json\n'
        'git commit -m "feat(sync): conflict resolution modal UI with single-field and multi-field support [Phase 2B, Prompt 7]"'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # PROMPT 8
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'PROMPT 8: Show Conflict Indicators in the Grid', level=1)

    add_gray_box(doc,
        'PROMPT 8 — What This Step Does',
        'This builds the visual indicators that show conflicts directly in the data grid: a small amber triangle in the corner of conflicted cells, an amber badge on conflicted rows, and a conflict count badge in the toolbar. These give Managers a quick way to spot and navigate to conflicts without opening a separate screen.',
        'Claude Code will create cell indicator, row badge, and toolbar badge components, plus the data layer to load conflict state. Component tests will verify rendering.',
        '8–10 minutes.'
    )

    add_blue_box(doc,
        'Read docs/reference/sync-engine.md lines 615–655 for the grid-level conflict rendering spec.\n'
        '\n'
        'Build the visual conflict indicators for the grid and the data layer to populate the _conflicts map.\n'
        '\n'
        'Target files:\n'
        '- apps/web/src/components/sync/CellConflictIndicator.tsx (new)\n'
        '- apps/web/src/components/sync/ConflictToolbarBadge.tsx (new)\n'
        '- apps/web/src/data/sync-conflicts.ts (new)\n'
        '\n'
        'What to build:\n'
        '\n'
        '1. Data functions in sync-conflicts.ts:\n'
        '   - getPendingConflictsForTable(tenantId, tableId) returns { [recordId]: { [fieldId]: ConflictMeta } }\n'
        '   - getPendingConflictCount(tenantId, tableId) returns the total count\n'
        '   - ConflictMeta type: { id, localValue, remoteValue, platform, createdAt }\n'
        '\n'
        '2. CellConflictIndicator.tsx:\n'
        '   - Renders 4px amber triangle in top-right corner (CSS :after or positioned SVG)\n'
        '   - 1px amber dashed underline on cell text\n'
        '   - Tooltip on hover: "Conflict: edited both locally and on {Platform}. Click to resolve."\n'
        '   - Click opens ConflictResolutionModal for this record\n'
        '   - Only renders when record._conflicts?.[fieldId] exists\n'
        '\n'
        '3. RowConflictBadge component:\n'
        '   - Shows amber warning badge in row number area when any cell in that row is conflicted\n'
        '   - Click opens ConflictResolutionModal showing all conflicts for that record\n'
        '\n'
        '4. ConflictToolbarBadge.tsx:\n'
        '   - Shows "{count} conflicts" in amber text\n'
        '   - Click filters grid to only conflicted records\n'
        '   - Hidden when count is 0\n'
        '\n'
        '5. CellWrapper component (export for Phase 3 grid integration):\n'
        '   - Wraps any cell with optional conflict indicator overlay\n'
        '   - Phase 3 TanStack Table column definitions will use this\n'
        '\n'
        'Acceptance criteria:\n'
        '- getPendingConflictsForTable() returns correct nested map\n'
        '- getPendingConflictCount() returns correct count\n'
        '- CellConflictIndicator renders 4px amber triangle + dashed underline\n'
        '- Tooltip shows conflict description with platform name\n'
        '- Click opens resolution modal\n'
        '- RowConflictBadge shows on rows with conflicts\n'
        '- ConflictToolbarBadge shows count, hidden when zero, click filters\n'
        '- CellWrapper exported for Phase 3\n'
        '- All text uses i18n\n'
        '- testTenantIsolation() passes for data queries\n'
        '- Component tests verify rendering\n'
        '- ESLint and TypeScript compile with zero errors\n'
        '- Coverage >= 80% on new files\n'
        '\n'
        'Do NOT build: full TanStack Table grid (Phase 3), cell renderers per field type (Phase 3), real-time push (Prompt 9), resolution action (Prompt 10), mobile indicators (Phase 3H).'
    )

    add_orange_box(doc,
        'Checkpoint — Verify Prompt 8\n'
        'Look for:\n'
        '- New files: CellConflictIndicator.tsx, ConflictToolbarBadge.tsx, sync-conflicts.ts\n'
        '- CellWrapper component exported\n'
        '- Component tests passing\n'
        '- No TypeScript or ESLint errors'
    )

    add_green_box(doc,
        'git add apps/web/src/components/sync/ apps/web/src/data/sync-conflicts.ts\n'
        'git commit -m "feat(sync): grid-level conflict rendering with cell indicators, row badges, and toolbar badge [Phase 2B, Prompt 8]"'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # INTEGRATION CHECKPOINT 2
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'INTEGRATION CHECKPOINT 2', level=1)

    add_gray_box(doc,
        'Integration Checkpoint 2 — Verify Prompts 5–8',
        'This verifies that the conflict detection system (Prompts 5–8) works correctly and integrates with the outbound sync infrastructure from Prompts 1–4.',
        'You will run the full test suite again. Everything should pass. Claude Code will also verify the conflict flow manually.',
        '3–5 minutes.'
    )

    add_blue_box(doc,
        'Run the full verification suite for Prompts 5–8:\n'
        '\n'
        'pnpm turbo typecheck\n'
        'pnpm turbo lint\n'
        'pnpm turbo test\n'
        'pnpm turbo test -- --coverage\n'
        '\n'
        'Then verify:\n'
        '- Three-way conflict detection correctly identifies conflicts vs clean changes\n'
        '- Last-write-wins auto-resolution applies remote values and preserves local in sync_conflicts\n'
        '- Manual mode preserves local canonical_data when conflicts detected\n'
        '- ConflictResolutionModal renders correctly with mock data (single + multi-field)\n'
        '- CellConflictIndicator renders amber triangle and tooltip\n'
        '- ConflictToolbarBadge shows count and hides when zero\n'
        '\n'
        'Fix any failures before proceeding.'
    )

    add_orange_box(doc,
        'Checkpoint — All Must Pass\n'
        '- TypeScript: zero errors\n'
        '- ESLint: zero errors\n'
        '- Tests: all green\n'
        '- Coverage: thresholds met\n'
        '\n'
        'If anything fails, paste the error into Claude Code and say "fix this". Do NOT proceed until everything passes.'
    )

    add_green_box(doc,
        'git add -A\n'
        'git commit -m "chore(verify): integration checkpoint 2 — conflict detection, resolution mode, UI components [Phase 2B, CP-2]"\n'
        'git push origin feat/phase-2b-outbound-sync-conflicts'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # PROMPT 9
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'PROMPT 9: Real-Time Conflict Updates', level=1)

    add_gray_box(doc,
        'PROMPT 9 — What This Step Does',
        'This wires up real-time events so that when a conflict is detected (or resolved), every user looking at that table sees the amber indicator appear (or disappear) instantly — no page refresh needed. It uses the Socket.io/Redis infrastructure built in Phase 1G to broadcast conflict events to all connected clients.',
        'Claude Code will create event definitions, update the sync worker to emit events, and build a client-side Zustand store for conflict state. Tests will verify the event pipeline.',
        '8–10 minutes.'
    )

    add_blue_box(doc,
        'Read docs/reference/sync-engine.md lines 694–731 for the real-time conflict push spec.\n'
        '\n'
        'Build the real-time event pipeline for conflict detection and resolution.\n'
        '\n'
        'Target files:\n'
        '- apps/realtime/src/events/sync-events.ts (new)\n'
        '- apps/worker/src/jobs/sync-inbound.ts (extend — emit event after conflict write)\n'
        '- apps/web/src/lib/sync-conflict-store.ts (new)\n'
        '\n'
        'What to build:\n'
        '\n'
        '1. Define events in sync-events.ts:\n'
        '   - SyncConflictDetectedEvent: { type, recordId, fieldId, conflictId, localValue, remoteValue, platform }\n'
        '   - SyncConflictResolvedEvent: { type, recordId, fieldId, conflictId, resolvedValue, resolution }\n'
        '\n'
        '2. After writeConflictRecords() in the inbound sync worker, emit sync.conflict_detected via Redis pub-sub:\n'
        '   - Channel: t:{tenantId}:table:{tableId}\n'
        '   - Socket.io broadcasts to all clients in the table room\n'
        '\n'
        '3. Build sync-conflict-store.ts — Zustand store:\n'
        '   - conflicts: { [recordId]: { [fieldId]: ConflictMeta } }\n'
        '   - setInitialConflicts(), addConflict(), removeConflict()\n'
        '   - getConflictsForRecord(), conflictCount()\n'
        '\n'
        '4. Wire Socket.io listeners on table room join:\n'
        '   - sync.conflict_detected → addConflict() → cell re-renders with indicator\n'
        '   - sync.conflict_resolved → removeConflict() → indicator disappears\n'
        '\n'
        '5. On table navigation: fetch initial conflicts via getPendingConflictsForTable() → setInitialConflicts()\n'
        '\n'
        'Acceptance criteria:\n'
        '- sync.conflict_detected emitted after conflict records written\n'
        '- sync.conflict_resolved event defined (emission wired in Prompt 10)\n'
        '- Socket.io broadcasts both events to table room clients\n'
        '- ConflictStore add/remove triggers cell re-render\n'
        '- Initial conflicts loaded on table navigation\n'
        '- Unit tests: store add, remove, initialize, count\n'
        '- Integration test: Socket.io event delivery\n'
        '- ESLint and TypeScript compile with zero errors\n'
        '- Coverage >= 80% on new files\n'
        '\n'
        'Do NOT build: conflict resolution Server Action (Prompt 10), record.updated events (Phase 3), room management changes (use Phase 1G), push notifications (Phase 2C).'
    )

    add_orange_box(doc,
        'Checkpoint — Verify Prompt 9\n'
        'Look for:\n'
        '- New files: sync-events.ts, sync-conflict-store.ts\n'
        '- Updated sync-inbound.ts with event emission\n'
        '- Zustand store tests passing\n'
        '- Socket.io integration test passing'
    )

    add_green_box(doc,
        'git add apps/realtime/src/events/sync-events.ts apps/worker/src/jobs/sync-inbound.ts apps/web/src/lib/sync-conflict-store.ts\n'
        'git commit -m "feat(sync): real-time conflict push via Socket.io/Redis with client-side _conflicts map [Phase 2B, Prompt 9]"'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # PROMPT 10
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'PROMPT 10: Wire Up Conflict Resolution with Undo', level=1)

    add_gray_box(doc,
        'PROMPT 10 — What This Step Does',
        'This connects the conflict resolution popup (Prompt 7) to a Server Action that actually saves the resolution decision. When a Manager clicks "Keep EveryStack", the conflict is resolved, the correct value is applied, and if needed, an outbound sync pushes it to the platform. There\'s also an 8-second undo window — if you change your mind, click Undo on the toast that appears.',
        'Claude Code will create the resolution Server Action, an undo mechanism, and a toast component. It will also wire the modal buttons to these actions. Tests will cover all resolution types and the undo flow.',
        '10–15 minutes.'
    )

    add_blue_box(doc,
        'Read docs/reference/sync-engine.md lines 603–608 and lines 717–745 for resolution actions and undo spec.\n'
        '\n'
        'Build the conflict resolution Server Action with optimistic UI, undo toast, and outbound sync.\n'
        '\n'
        'Target files:\n'
        '- apps/web/src/actions/sync-conflict-resolve.ts (new)\n'
        '- apps/web/src/components/sync/ConflictResolutionModal.tsx (extend)\n'
        '- apps/web/src/components/sync/UndoResolveToast.tsx (new)\n'
        '\n'
        'What to build:\n'
        '\n'
        '1. resolveConflict Server Action:\n'
        '   - Input: { conflictId, resolution: resolved_local | resolved_remote | resolved_merged, mergedValue? }\n'
        '   - Validates via Zod, verifies conflict exists and is pending\n'
        '   - Permission: Owner/Admin/Manager only\n'
        '   - In single transaction: update sync_conflicts status/resolved_by/resolved_at, update canonical_data with resolved value, update search_vector, update sync_metadata\n'
        '   - Emit sync.conflict_resolved via Redis pub-sub\n'
        '   - Emit record.updated event\n'
        '   - resolved_local or resolved_merged: call enqueueOutboundSync() with P1 priority\n'
        '   - resolved_remote: no outbound sync needed\n'
        '   - Generate undo token (UUIDv7), cache previous state in Redis with 8-second TTL\n'
        '   - Return { success, undoToken }\n'
        '\n'
        '2. undoConflictResolution(undoToken) Server Action:\n'
        '   - If Redis key exists (within 8s): revert sync_conflicts to pending, restore canonical_data, cancel outbound job, emit sync.conflict_detected, delete Redis key\n'
        '   - If expired: return { success: false }\n'
        '\n'
        '3. UndoResolveToast.tsx (shadcn/ui Toast):\n'
        '   - Shows "Conflict resolved. [Undo]"\n'
        '   - Auto-dismisses after 8 seconds\n'
        '   - Undo button calls undoConflictResolution()\n'
        '   - Success: "Resolution undone." / Failure: "Undo no longer available."\n'
        '\n'
        '4. Wire ConflictResolutionModal buttons:\n'
        '   - Keep EveryStack → resolveConflict({ resolution: resolved_local })\n'
        '   - Keep {Platform} → resolveConflict({ resolution: resolved_remote })\n'
        '   - Edit → inline editor → resolveConflict({ resolution: resolved_merged, mergedValue })\n'
        '   - On success: removeConflict() on store (optimistic), show UndoResolveToast\n'
        '\n'
        'Acceptance criteria:\n'
        '- resolveConflict() updates sync_conflicts, canonical_data, search_vector, sync_metadata in single transaction\n'
        '- Permission: Owner/Admin/Manager only (403 for others)\n'
        '- resolved_local/merged triggers enqueueOutboundSync() P1\n'
        '- resolved_remote does NOT trigger outbound\n'
        '- sync.conflict_resolved event emitted\n'
        '- record.updated event emitted\n'
        '- Undo works within 8 seconds\n'
        '- Undo fails gracefully after 8 seconds\n'
        '- Toast renders and auto-dismisses\n'
        '- Modal buttons wired, indicator disappears optimistically\n'
        '- testTenantIsolation() passes\n'
        '- Unit tests: all 3 resolution types, undo in/out of window, permission denial\n'
        '- ESLint and TypeScript compile with zero errors\n'
        '- Coverage >= 90% on sync-conflict-resolve.ts\n'
        '\n'
        'Do NOT build: bulk resolution (Prompt 11), role-based visibility (Prompt 11), audit log (Prompt 11), feature interaction rules (Prompt 11).'
    )

    add_orange_box(doc,
        'Checkpoint — Verify Prompt 10\n'
        'Look for:\n'
        '- New files: sync-conflict-resolve.ts, UndoResolveToast.tsx\n'
        '- Updated ConflictResolutionModal.tsx with wired buttons\n'
        '- Resolution + undo tests passing\n'
        '- No TypeScript or ESLint errors'
    )

    add_green_box(doc,
        'git add apps/web/src/actions/sync-conflict-resolve.ts apps/web/src/components/sync/UndoResolveToast.tsx apps/web/src/components/sync/ConflictResolutionModal.tsx\n'
        'git commit -m "feat(sync): conflict resolution Server Action with optimistic UI, undo toast, and outbound sync [Phase 2B, Prompt 10]"'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # PROMPT 11
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'PROMPT 11: Roles, Bulk Resolution, Audit Trail, and Feature Rules', level=1)

    add_gray_box(doc,
        'PROMPT 11 — What This Step Does',
        'This is the final build prompt. It completes the conflict system with four things: (1) Role-based visibility — Managers see full controls, Team Members see a read-only indicator, Viewers see nothing. (2) Bulk resolution — resolve all conflicts at once with one click. (3) Audit trail — every resolution is logged for accountability. (4) Feature interaction rules — how conflicts interact with automations, cross-links, portals, and search.',
        'Claude Code will update existing components with role logic, add bulk resolution Server Actions, add audit log calls, and create a feature interaction utility file. This is the largest prompt in the phase.',
        '10–15 minutes.'
    )

    add_blue_box(doc,
        'Read docs/reference/sync-engine.md lines 683–693 (role visibility), lines 609 (bulk resolution), and lines 757–793 (feature interactions and audit trail).\n'
        '\n'
        'Complete the conflict system with role enforcement, bulk actions, audit logging, and feature interaction rules.\n'
        '\n'
        'Target files:\n'
        '- apps/web/src/components/sync/CellConflictIndicator.tsx (extend)\n'
        '- apps/web/src/components/sync/ConflictToolbarBadge.tsx (extend)\n'
        '- apps/web/src/actions/sync-conflict-resolve.ts (extend)\n'
        '- apps/web/src/components/sync/ConflictResolutionModal.tsx (extend)\n'
        '- packages/shared/sync/conflict-interactions.ts (new)\n'
        '\n'
        'What to build:\n'
        '\n'
        '1. Role-based conflict visibility:\n'
        '   - Owner/Admin: full indicators + resolution modal\n'
        '   - Manager: full indicators + resolution modal (on permitted bases)\n'
        '   - Team Member: amber indicator visible, but click shows read-only tooltip: "This field has a sync conflict. A Manager will resolve it." NO resolution modal.\n'
        '   - Viewer: NO indicators at all\n'
        '   - Update CellConflictIndicator and ConflictToolbarBadge accordingly\n'
        '\n'
        '2. Bulk conflict resolution Server Actions:\n'
        '   - bulkResolveConflicts({ recordId, resolution }) — resolves ALL pending conflicts on one record\n'
        '   - bulkResolveTableConflicts({ tableId, resolution }) — resolves ALL pending conflicts in a table\n'
        '   - Both emit sync.conflict_resolved events per conflict\n'
        '   - Both generate a single undo token for the entire batch (8-second Redis TTL)\n'
        '   - Wire to modal "Keep All" buttons\n'
        '\n'
        '3. Conflict audit trail — after every resolution (including bulk):\n'
        '   - Call writeAuditLog() with action: \'sync_conflict.resolved\'\n'
        '   - Details: conflictId, fieldId, resolution, previousValue, resolvedValue, platform\n'
        '   - Activity tab format: "{User} resolved sync conflict on \'{fieldName}\' — kept EveryStack value..."\n'
        '   - Include in resolveConflict() and bulkResolveConflicts() in same transaction\n'
        '\n'
        '4. Feature interaction rules (conflict-interactions.ts):\n'
        '   - shouldRecomputeOnResolution(fieldId, tableFields) returns { formulaFields, crossLinkFields, requiresTsvectorUpdate }\n'
        '   - Document with code comments: automations run with current value, cross-links show current value, portals hide conflicts, formulas compute with current value, search uses current value\n'
        '\n'
        'Acceptance criteria:\n'
        '- Owner/Admin/Manager see full conflict indicators + modal\n'
        '- Team Member sees indicator but CANNOT open modal (read-only tooltip)\n'
        '- Viewer sees NO indicators\n'
        '- bulkResolveConflicts() resolves all on a record\n'
        '- bulkResolveTableConflicts() resolves all in a table\n'
        '- Bulk generates single undo token\n'
        '- writeAuditLog() called for every resolution\n'
        '- Audit message matches format spec\n'
        '- shouldRecomputeOnResolution() identifies downstream fields\n'
        '- testTenantIsolation() passes\n'
        '- Unit tests: role visibility, bulk resolve, audit log, feature interactions\n'
        '- ESLint and TypeScript compile with zero errors\n'
        '- Coverage >= 80% on all new/modified files\n'
        '\n'
        'Do NOT build: automation trigger integration (Phase 4), cross-link cascade (Phase 3B), portal cache invalidation (Phase 3E), formula recalculation (post-MVP), mobile layout (Phase 3H).'
    )

    add_orange_box(doc,
        'Checkpoint — Verify Prompt 11\n'
        'Look for:\n'
        '- Updated indicator components with role logic\n'
        '- Bulk resolution actions added to sync-conflict-resolve.ts\n'
        '- Audit log calls in resolution actions\n'
        '- New file: conflict-interactions.ts\n'
        '- All tests passing'
    )

    add_green_box(doc,
        'git add apps/web/src/components/sync/ apps/web/src/actions/sync-conflict-resolve.ts packages/shared/sync/conflict-interactions.ts\n'
        'git commit -m "feat(sync): role-based conflict visibility, bulk resolution, audit trail, and feature interaction rules [Phase 2B, Prompt 11]"'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # FINAL INTEGRATION CHECKPOINT
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'FINAL INTEGRATION CHECKPOINT', level=1)

    add_gray_box(doc,
        'Final Checkpoint — Full Phase 2B Verification',
        'This is the final verification for the entire phase. You are running the complete test suite, verifying coverage thresholds, and checking the end-to-end conflict flow. Once everything passes, you will push and open a Pull Request to merge Phase 2B into main.',
        'You will run several checks and then Claude Code will run through the end-to-end flow manually. Everything must pass before opening the PR.',
        '5–10 minutes.'
    )

    add_blue_box(doc,
        'Run the complete Phase 2B verification:\n'
        '\n'
        'pnpm turbo typecheck\n'
        'pnpm turbo lint\n'
        'pnpm turbo test\n'
        'pnpm turbo test -- --coverage\n'
        'pnpm turbo check:i18n\n'
        '\n'
        'Verify coverage thresholds:\n'
        '- packages/shared/sync/ >= 90% lines, 85% branches\n'
        '- apps/web/src/actions/ >= 90% lines, 85% branches\n'
        '- apps/worker/src/jobs/ >= 85% lines, 80% branches\n'
        '\n'
        'Then run the end-to-end verification:\n'
        '- Simulate inbound sync with conflicting values → verify sync_conflicts created\n'
        '- Verify real-time event arrives on client → amber indicator renders\n'
        '- Open resolution modal → resolve as "Keep EveryStack"\n'
        '- Verify outbound sync enqueued → amber indicator disappears\n'
        '- Verify undo toast appears → undo works within 8 seconds\n'
        '- Verify role visibility: Team Member sees indicator but cannot resolve; Viewer sees nothing\n'
        '- Bulk resolve multiple conflicts → verify single undo token\n'
        '- Verify audit log entries with correct format\n'
        '- Verify EXPLAIN shows index scan for canonicalFieldExpression() query\n'
        '\n'
        'Fix any failures before proceeding to the PR.'
    )

    add_orange_box(doc,
        'Final Checkpoint — ALL Must Pass\n'
        '- TypeScript: zero errors\n'
        '- ESLint: zero errors\n'
        '- i18n: no hardcoded English strings\n'
        '- Tests: all green\n'
        '- Coverage: all thresholds met\n'
        '- End-to-end conflict flow: working\n'
        '- Role visibility: correct\n'
        '- Audit trail: entries present\n'
        '\n'
        'If ANYTHING fails, paste the error into Claude Code and say "fix this".\n'
        'Do NOT open the PR until every check passes.'
    )

    add_green_box(doc,
        'git add -A\n'
        'git commit -m "chore(verify): final integration checkpoint — complete Phase 2B verification [Phase 2B, CP-3]"\n'
        'git push origin feat/phase-2b-outbound-sync-conflicts'
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # OPEN PR
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, 'OPEN PULL REQUEST', level=1)

    add_gray_box(doc,
        'Pull Request — Merge Phase 2B to Main',
        'This is the final step. You are opening a Pull Request on GitHub to merge all of Phase 2B into the main branch. Once the PR is reviewed and all CI checks pass, you can merge it.',
        'A PR will be created on GitHub. You can review it in your browser.',
        '1–2 minutes.'
    )

    add_green_box(doc,
        'gh pr create --title "Phase 2B — Synced Data Performance, Outbound Sync, Conflict Resolution" --body "## Summary\n'
        '- JSONB expression indexes for grid query performance (<200ms on 50K records)\n'
        '- sync_metadata tracking with last_synced_values per field\n'
        '- Outbound sync pipeline (BullMQ + fromCanonical + rate limiting)\n'
        '- Optimistic cell edit with outbound sync enqueue\n'
        '- Three-way conflict detection on inbound sync\n'
        '- Last-write-wins default + manual resolution mode toggle\n'
        '- Conflict resolution modal (single + multi-field)\n'
        '- Grid-level conflict indicators (cell triangle, row badge, toolbar badge)\n'
        '- Real-time conflict push via Socket.io/Redis\n'
        '- Resolution Server Action with 8-second undo\n'
        '- Role-based conflict visibility\n'
        '- Bulk conflict resolution\n'
        '- Conflict audit trail\n'
        '- Feature interaction rules\n'
        '\n'
        '## Test Plan\n'
        '- [ ] All unit tests pass\n'
        '- [ ] All integration tests pass\n'
        '- [ ] Coverage thresholds met\n'
        '- [ ] End-to-end conflict flow verified\n'
        '- [ ] Role visibility verified\n'
        '- [ ] i18n check passes"'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Phase 2B is complete. Next up: Phase 2C (Smart Polling, Notion Adapter, Sync Settings Dashboard).')
    run.bold = True
    run.font.size = Pt(12)

    # ── Save ──
    output_path = '/Users/stevenyeats/Documents/EveryStack/docs/Playbooks/Phase-2B-Operator-Guide.docx'
    doc.save(output_path)
    print(f'Saved to: {output_path}')


if __name__ == '__main__':
    main()
