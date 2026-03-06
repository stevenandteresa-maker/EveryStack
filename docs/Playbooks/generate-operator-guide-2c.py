#!/usr/bin/env python3
"""Generate Phase 2C Operator Guide as a color-coded .docx file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Style definitions ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Color constants
BLUE_BG = "D6E4F0"      # Prompt containers
GREEN_BG = "D9EAD3"     # Git containers
ORANGE_BG = "FCE5CD"    # Checkpoint/verification containers
GRAY_BG = "EFEFEF"      # Context/explanation containers
WHITE = "FFFFFF"
DARK_TEXT = "333333"
HEADER_BLUE = "1F4E79"
MONO_FONT = "Consolas"


def add_colored_container(doc, label, content_lines, bg_color, mono=True):
    """Add a color-coded container with a label and content."""
    # Create a single-cell table for the container
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set table width
    table.columns[0].width = Inches(6.5)

    cell = table.cell(0, 0)

    # Set cell background color
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

    # Add cell borders
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(borders)

    # Add label
    label_para = cell.paragraphs[0]
    label_run = label_para.add_run(label)
    label_run.bold = True
    label_run.font.size = Pt(8)
    label_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    label_run.font.name = 'Calibri'
    label_para.space_after = Pt(4)
    label_para.space_before = Pt(2)

    # Add content
    for line in content_lines:
        para = cell.add_paragraph()
        run = para.add_run(line)
        if mono:
            run.font.name = MONO_FONT
            run.font.size = Pt(9)
        else:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        para.space_after = Pt(1)
        para.space_before = Pt(1)

    doc.add_paragraph()  # spacer


def add_context_block(doc, what_builds, what_see, how_long):
    """Add a gray context explanation block."""
    lines = []
    lines.append(f"What This Builds: {what_builds}")
    lines.append("")
    lines.append(f"What You'll See When It's Done: {what_see}")
    lines.append("")
    lines.append(f"How Long This Typically Takes: {how_long}")
    add_colored_container(doc, "CONTEXT — READ FOR UNDERSTANDING", lines, GRAY_BG, mono=False)


def add_heading(doc, text, level=1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return heading


def add_separator(doc):
    """Add a visual separator."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)


def add_checkpoint_note(doc, text):
    """Add a checkpoint/verification note after a container."""
    para = doc.add_paragraph()
    run = para.add_run(f"Checkpoint: {text}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)


def add_body(doc, text):
    """Add a regular body paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return para


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading("Phase 2C — Operator Guide", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.size = Pt(28)

subtitle = doc.add_heading("Notion Adapter, Error Recovery, Sync Dashboard", level=2)
for run in subtitle.runs:
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
add_body(doc, "For: Steven (non-technical founder)")
add_body(doc, "Format: Follow this guide top-to-bottom. Copy containers exactly. Do not skip steps.")
doc.add_paragraph()

# Color key table
add_heading(doc, "Color Key", level=2)
color_table = doc.add_table(rows=5, cols=2)
color_table.style = 'Table Grid'
headers = [("Color", "Meaning")]
rows_data = [
    ("Blue", "Prompt to paste into Claude Code"),
    ("Green", "Git command to execute"),
    ("Orange / Amber", "Checkpoint — stop and verify"),
    ("Gray", "Context — read for understanding"),
]
# Header row
for i, (h1, h2) in enumerate([("Color", "Meaning")]):
    cell0 = color_table.cell(0, 0)
    cell1 = color_table.cell(0, 1)
    cell0.text = h1
    cell1.text = h2
    for cell in [cell0, cell1]:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

bg_colors = [BLUE_BG, GREEN_BG, ORANGE_BG, GRAY_BG]
for idx, (color_name, meaning) in enumerate(rows_data):
    row_idx = idx + 1
    cell0 = color_table.cell(row_idx, 0)
    cell1 = color_table.cell(row_idx, 1)
    cell0.text = color_name
    cell1.text = meaning
    # Color the first cell
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_colors[idx]}"/>')
    cell0._tc.get_or_add_tcPr().append(shading)

doc.add_page_break()

# ============================================================
# SETUP
# ============================================================
add_heading(doc, "SETUP", level=1)

add_colored_container(doc, "CONTEXT — READ FOR UNDERSTANDING", [
    "What this phase builds: Phase 2C adds three major capabilities to EveryStack's sync engine.",
    "",
    "First, it connects to Notion — so your users can sync data from Notion databases just",
    "like they already can with Airtable.",
    "",
    "Second, it builds error recovery — when things go wrong with a sync (expired login,",
    "platform outage, record failures), the system detects the problem, tells the user what",
    "happened, and offers clear resolution steps.",
    "",
    "Third, it builds the Sync Settings Dashboard — a 6-tab management panel where workspace",
    "managers can monitor, troubleshoot, and control their sync connections.",
    "",
    "Before you start, make sure:",
    "  - Your terminal is open and you're in the EveryStack monorepo",
    "  - Docker Desktop is running (PostgreSQL, Redis, PgBouncer need to be up)",
    "  - You're on the main branch with all Phase 2A and 2B work merged",
    "  - Claude Code is open (VS Code or terminal)",
], GRAY_BG, mono=False)

add_heading(doc, "Step 1: Open Your Terminal", level=2)

add_colored_container(doc, "GIT COMMAND", [
    "cd ~/Documents/EveryStack",
    "git checkout main",
    "git pull origin main",
], GREEN_BG)

add_body(doc, "Confirm you see no errors. You should be on main with the latest code.")

add_heading(doc, "Step 2: Load the Skill Files", level=2)
add_body(doc, "Before any prompts, paste this into Claude Code so it knows the project conventions:")

add_colored_container(doc, "PASTE INTO CLAUDE CODE", [
    "Read these skill files to understand project conventions before we begin building:",
    "- docs/skills/backend/SKILL.md",
    "- docs/skills/ux-ui/SKILL.md",
    "- docs/skills/phase-context/SKILL.md",
], BLUE_BG)

add_body(doc, "Wait for Claude Code to confirm it has read the files. This takes about 30 seconds.")

add_separator(doc)
doc.add_page_break()

# ============================================================
# PROMPT 1
# ============================================================
add_heading(doc, "PROMPT 1: Build the Notion Data Translator (Inbound)", level=1)

add_context_block(doc,
    "Right now, EveryStack can read data from Airtable and translate it into its own internal format. This prompt does the same thing for Notion. Notion stores data differently — it uses \"pages\" with \"properties\" instead of \"rows\" with \"fields\" — so we need a translator that converts Notion's format into EveryStack's universal format. Think of it like hiring a translator who speaks Notion and can write everything down in EveryStack's language.",
    "Claude Code will create 3 new files in the sync adapters folder. It will register about 20 Notion field types (like text, numbers, dates, checkboxes, etc.) in the system. Tests will run and you should see all green checkmarks covering every field type.",
    "5–10 minutes."
)

prompt1_lines = [
    "We are starting Phase 2C. First, create and checkout a new branch:",
    "git checkout -b feat/phase-2c-notion-error-recovery-dashboard",
    "",
    "Then build Prompt 1: NotionAdapter.toCanonical() for all MVP field types.",
    "",
    "Load context from: sync-engine.md lines 30-87 (Core Pattern, Source References,",
    "Field Type Registry), sync-engine.md lines 93-106 (Sync Setup - Step 2 databases",
    "for Notion), data-model.md lines 150-200 (Field Type Canonical Shapes).",
    "",
    "Target files:",
    "- packages/shared/sync/adapters/notion-adapter.ts",
    "- packages/shared/sync/adapters/notion-field-transforms.ts",
    "- packages/shared/sync/adapters/notion-types.ts",
    "",
    "No migration required.",
    "",
    "TASK:",
    "Build the NotionAdapter.toCanonical() implementation, mapping Notion's page property",
    "model to EveryStack's canonical JSONB shape. This is the inbound direction — Notion",
    "API response to canonical JSONB.",
    "",
    "1. Create packages/shared/sync/adapters/notion-types.ts:",
    "Define TypeScript types for Notion API responses. The Notion API returns page objects",
    "with a properties map where each property has a type field and type-specific value",
    "structure. Key types to model:",
    "",
    "type NotionPropertyType =",
    "  | 'title' | 'rich_text' | 'number' | 'select' | 'multi_select'",
    "  | 'date' | 'people' | 'files' | 'checkbox' | 'url' | 'email'",
    "  | 'phone_number' | 'formula' | 'relation' | 'rollup'",
    "  | 'created_time' | 'created_by' | 'last_edited_time' | 'last_edited_by'",
    "  | 'status' | 'unique_id';",
    "",
    "2. Create packages/shared/sync/adapters/notion-field-transforms.ts:",
    "Implement per-field-type transform functions following the same pattern as Airtable's",
    "transforms. Key mappings:",
    "- title -> text (extract plain text, preserve rich text in source_refs)",
    "- rich_text -> long_text (extract plain text, preserve rich text array in source_refs)",
    "- number -> number (direct)",
    "- select -> single_select (map {id, name, color} with source_refs)",
    "- multi_select -> multi_select (array of options with source_refs per option)",
    "- date -> date (map {start, end, time_zone} to canonical date shape)",
    "- people -> user (Notion user objects to canonical user references)",
    "- files -> attachment (file objects to canonical attachment shape)",
    "- checkbox -> checkbox (direct boolean)",
    "- url -> url, email -> email, phone_number -> phone (direct strings)",
    "- formula -> formula (read-only, isLossless: false, extract computed value)",
    "- relation -> link (Notion relation IDs to cross-link refs, store page IDs in source_refs)",
    "- rollup -> rollup (read-only, isLossless: false, extract computed value)",
    "- created_time -> created_time, last_edited_time -> last_edited_time (ISO 8601)",
    "- created_by -> created_by, last_edited_by -> last_edited_by (user reference)",
    "- status -> single_select (preserve group info in source_refs)",
    "- unique_id -> auto_number (read-only counter)",
    "",
    "3. Register all Notion field transforms in the FieldTypeRegistry:",
    "For each mapping, call FieldTypeRegistry.register() with platform 'notion'. Follow the",
    "exact registration pattern used for Airtable in Phase 2A. Mark lossy fields: formula,",
    "rollup, created_time, created_by, last_edited_time, last_edited_by, unique_id ->",
    "isLossless: false.",
    "",
    "4. Create packages/shared/sync/adapters/notion-adapter.ts:",
    "Implement NotionAdapter class implementing the PlatformAdapter interface. For toCanonical():",
    "- Accept a Notion page object",
    "- Iterate over its properties map",
    "- For each property, look up the registered transform in FieldTypeRegistry",
    "  for ('notion', notionPropertyType)",
    "- Call the transform's toCanonical() to produce the canonical field value",
    "- Assemble the canonical JSONB keyed by EveryStack fields.id",
    "  (looked up via synced_field_mappings)",
    "- Populate source_refs with Notion-specific identifiers (page ID, property IDs, option IDs)",
    "",
    "5. Handle Notion's rich text model:",
    "Notion's title and rich_text properties return arrays of rich text objects with annotations",
    "(bold, italic, code, color, links). The canonical text/long_text types store plain text",
    "in value. Preserve the full rich text array in source_refs.notion.rich_text for lossless",
    "round-tripping on fromCanonical().",
    "",
    "ACCEPTANCE CRITERIA:",
    "- NotionAdapter class implements the PlatformAdapter interface",
    "- toCanonical() correctly transforms all ~20 Notion property types to canonical JSONB shapes",
    "- Each Notion property type has a corresponding FieldTypeRegistry.register() call",
    "  with platform 'notion'",
    "- Lossy fields (formula, rollup, created_time, created_by, last_edited_time,",
    "  last_edited_by, unique_id) are marked isLossless: false",
    "- source_refs correctly preserves Notion-specific identifiers",
    "- Rich text properties extract plain text for value and preserve full rich text in source_refs",
    "- select/multi_select/status transforms preserve Notion option IDs in source_refs",
    "- relation property maps to cross-link references with Notion page IDs in source_refs",
    "- Unit tests cover all ~20 field type transforms with representative Notion API fixtures",
    "- Edge cases tested: null/empty properties, rich text with mixed annotations,",
    "  dates with/without end dates, empty relations",
    "- testTenantIsolation() passes for any new data access functions",
    "- ESLint and TypeScript compile with zero errors",
    "- Coverage >= 80% on new files",
    "",
    "DO NOT BUILD:",
    "- fromCanonical() direction (Prompt 2)",
    "- Notion filter pushdown (Prompt 2)",
    "- Notion OAuth flow integration (Prompt 2)",
    "- Notion API client or HTTP layer — use the @notionhq/client SDK",
    "- SmartSuite adapter (Phase 3)",
    "- Block-level content sync (only page-level properties)",
]
add_colored_container(doc, "PASTE INTO CLAUDE CODE", prompt1_lines, BLUE_BG)

add_body(doc, "Wait for Claude Code to finish. You should see 3 new files created and tests passing.")

add_colored_container(doc, "GIT COMMAND", [
    "git add packages/shared/sync/adapters/notion-adapter.ts \\",
    "       packages/shared/sync/adapters/notion-field-transforms.ts \\",
    "       packages/shared/sync/adapters/notion-types.ts",
    "git add -A packages/shared/sync/adapters/__tests__/",
    'git commit -m "feat(sync): implement NotionAdapter.toCanonical() for all MVP field types [Phase 2C, Prompt 1]"',
], GREEN_BG)

add_checkpoint_note(doc, "Confirm the commit succeeded with no errors. You see the 3 new adapter files in the commit.")

add_separator(doc)
doc.add_page_break()

# ============================================================
# PROMPT 2
# ============================================================
add_heading(doc, "PROMPT 2: Complete the Notion Adapter (Outbound + OAuth + Filters)", level=1)

add_context_block(doc,
    "Prompt 1 taught EveryStack how to read from Notion. This prompt teaches it how to write back to Notion. It also adds the ability for users to connect their Notion account (via OAuth login), adds smart filtering so EveryStack only pulls the records the user wants, and wires everything into the existing sync pipeline. After this, Notion sync is fully functional end-to-end.",
    "The sync setup wizard will show \"Notion\" as a platform option alongside Airtable. Claude Code will extend several existing files and create a few new ones. Integration tests will verify that a mock Notion sync works end-to-end.",
    "8–12 minutes."
)

prompt2_lines = [
    "Build Prompt 2: NotionAdapter.fromCanonical(), Notion Filter Pushdown, OAuth Flow,",
    "and Sync Pipeline Integration.",
    "",
    "Load context from: sync-engine.md lines 190-260 (Filter Pushdown, Estimated Record",
    "Count), sync-engine.md lines 93-106 (Sync Setup steps).",
    "",
    "Target files:",
    "- packages/shared/sync/adapters/notion-adapter.ts (extend)",
    "- packages/shared/sync/adapters/notion-field-transforms.ts (extend — add fromCanonical)",
    "- packages/shared/sync/adapters/notion-filter.ts (new)",
    "- apps/web/src/lib/notion-oauth.ts (new)",
    "- apps/worker/src/jobs/sync-inbound.ts (extend)",
    "",
    "No migration required.",
    "",
    "TASK:",
    "Complete the Notion adapter with outbound transforms, filter pushdown, OAuth integration,",
    "and full sync pipeline wiring.",
    "",
    "1. Implement fromCanonical() for all Notion field types:",
    "Extend notion-field-transforms.ts with fromCanonical() for each writable field type.",
    "Key outbound mappings:",
    "- text -> title rich text array (reconstruct from source_refs.notion.rich_text if",
    "  available, otherwise create plain text rich text object)",
    "- long_text -> rich_text array (same reconstruction logic)",
    "- number -> number value",
    "- single_select -> select object with {id} from source_refs or {name} for new options",
    "- multi_select -> array of {id} or {name} objects",
    "- date -> {start, end, time_zone} object",
    "- checkbox -> boolean",
    "- url, email, phone -> string values",
    "Read-only fields (formula, rollup, created_time, created_by, last_edited_time,",
    "last_edited_by, unique_id) must NOT be included in outbound updates.",
    "Register the fromCanonical direction for all writable Notion types in FieldTypeRegistry.",
    "",
    "2. Build packages/shared/sync/adapters/notion-filter.ts:",
    "Implement FilterRule[] to Notion filter JSON translator.",
    "export function translateToNotionFilter(",
    "  rules: FilterRule[], fieldMappings: SyncedFieldMapping[]",
    "): NotionFilter | undefined",
    "Map EveryStack filter operators to Notion filter conditions:",
    "- equals -> equals (text, number, select, checkbox)",
    "- not_equals -> does_not_equal",
    "- contains -> contains (text, rich_text)",
    "- not_contains -> does_not_contain",
    "- is_empty -> is_empty: true",
    "- is_not_empty -> is_not_empty: true",
    "- greater_than -> greater_than (number, date)",
    "- less_than -> less_than (number, date)",
    "- is_before -> before (date)",
    "- is_after -> after (date)",
    "Use synced_field_mappings to resolve EveryStack field IDs to Notion property IDs.",
    "",
    "3. Implement Notion OAuth flow in apps/web/src/lib/notion-oauth.ts:",
    "Follow the same OAuth pattern established for Airtable in Phase 2A:",
    "- Build the Notion OAuth authorization URL with required scopes",
    "- Handle the OAuth callback, exchange code for access token",
    "- Store the encrypted token in base_connections.credentials",
    "- Notion OAuth uses Authorization: Bearer {token} for all API calls",
    "",
    "4. Integrate Notion into the sync setup wizard:",
    "Extend the existing 3-step sync setup wizard to support Notion:",
    "- Step 1 (Authenticate): Add Notion as a platform choice.",
    "- Step 2 (Select Database): After auth, fetch databases via POST /v1/search",
    "  with filter: { property: 'object', value: 'database' }.",
    "- Step 3 (Select Tables & Filters): Same UI as Airtable.",
    "",
    "5. Wire NotionAdapter into the inbound sync pipeline:",
    "Update apps/worker/src/jobs/sync-inbound.ts to detect platform === 'notion' and use",
    "NotionAdapter instead of AirtableAdapter. Use POST /v1/databases/{id}/query with",
    "pagination (start_cursor). Apply filter pushdown. Process pages in batches.",
    "",
    "6. Register Notion rate limit:",
    "{ platform: 'notion', limits: [{ scope: 'per_integration', limit: 3, window: 1000 }] }",
    "",
    "7. Estimated record count for Notion:",
    "Use page_size=1 on filtered database query to check if results exist. Show as",
    "'~N records (estimated)' with tilde.",
    "",
    "ACCEPTANCE CRITERIA:",
    "- fromCanonical() correctly transforms canonical values back to Notion API update format",
    "- Read-only fields are skipped in fromCanonical()",
    "- translateToNotionFilter() correctly maps FilterRule[] operators to Notion filters",
    "- Filter translation handles compound rules (multiple conditions with and)",
    "- Notion OAuth flow completes end-to-end: auth URL -> callback -> token storage",
    "- Sync setup wizard shows Notion as a platform option",
    "- Inbound sync pipeline dispatches to NotionAdapter when platform === 'notion'",
    "- Notion inbound sync uses POST /v1/databases/{id}/query with pagination + filter",
    "- Notion rate limit (3 req/s per integration) is registered",
    "- Record count estimation works for Notion databases",
    "- source_refs round-tripping: toCanonical() -> fromCanonical() preserves identifiers",
    "- Integration test: mock Notion API, run full inbound sync, verify canonical_data shape",
    "- testTenantIsolation() passes for any new data access functions",
    "- ESLint and TypeScript compile with zero errors",
    "- Coverage >= 80% on new files",
    "",
    "DO NOT BUILD:",
    "- Smart polling or adaptive intervals (Prompt 3)",
    "- Priority-based scheduling (Prompt 4)",
    "- Error recovery flows (Prompts 6-9)",
    "- Sync dashboard (Prompt 10)",
    "- SmartSuite adapter (Phase 3)",
    "- Notion block-level content sync",
]
add_colored_container(doc, "PASTE INTO CLAUDE CODE", prompt2_lines, BLUE_BG)

add_body(doc, "Wait for Claude Code to finish. You should see the sync wizard now listing Notion alongside Airtable.")

add_colored_container(doc, "GIT COMMAND", [
    "git add packages/shared/sync/adapters/ apps/web/src/lib/notion-oauth.ts \\",
    "       apps/worker/src/jobs/sync-inbound.ts",
    "git add -A",
    'git commit -m "feat(sync): implement NotionAdapter.fromCanonical(), filter pushdown, OAuth, and sync pipeline integration [Phase 2C, Prompt 2]"',
], GREEN_BG)

add_checkpoint_note(doc, "Confirm the commit succeeded.")

add_separator(doc)
doc.add_page_break()

# ============================================================
# PROMPT 3
# ============================================================
add_heading(doc, "PROMPT 3: Make Sync Polling Smarter", level=1)

add_context_block(doc,
    "Until now, EveryStack polls external platforms on a fixed schedule. This prompt makes polling intelligent. If a user is actively looking at a table, it syncs every 30 seconds. If the table is open but not on screen, it syncs every 5 minutes. If nobody is using the workspace, it syncs every 30 minutes. It also teaches the system to skip tables that have been \"converted\" to native EveryStack tables. Think of it like a smart thermostat — it adjusts based on actual usage.",
    "A new scheduler file will be created. Tests will verify that different \"visibility states\" produce different polling intervals. The system will also register Airtable webhook listeners where available, so some syncs can be instant.",
    "5–8 minutes."
)

prompt3_lines = [
    "Build Prompt 3: Smart Polling with Adaptive Intervals and Converted Table Skip Logic.",
    "",
    "This prompt is independent of the Notion adapter prompts.",
    "",
    "Load context from: sync-engine.md lines 459-487 (Smart Polling subsection —",
    "adaptive intervals, visibility-based scheduling, webhook listeners).",
    "",
    "Target files:",
    "- apps/worker/src/jobs/sync-scheduler.ts (new or extend)",
    "- apps/worker/src/jobs/sync-inbound.ts (extend — converted table check)",
    "- packages/shared/sync/types.ts (extend — polling states)",
    "",
    "No migration required.",
    "",
    "TASK:",
    "Upgrade the sync scheduler from fixed-interval polling to adaptive intervals based on",
    "table visibility state, and add converted table skip logic.",
    "",
    "1. Define polling interval tiers:",
    "ACTIVE_VIEWING: 30_000 (30 seconds — user has this table open)",
    "TAB_OPEN_NOT_VISIBLE: 300_000 (5 minutes — workspace open, table not active)",
    "WORKSPACE_INACTIVE: 1_800_000 (30 minutes — workspace not accessed recently)",
    "EVENT_DRIVEN: null (Airtable webhooks — no polling, event-triggered)",
    "",
    "2. Build the adaptive scheduler in apps/worker/src/jobs/sync-scheduler.ts:",
    "export function getPollingInterval(",
    "  platform: string, tableVisibility: TableVisibility, hasWebhook: boolean",
    "): number | null",
    "Where TableVisibility is derived from real-time connection state:",
    "- 'active' — client has this table's Socket.io room joined",
    "- 'background' — client in same workspace, not viewing this table",
    "- 'inactive' — no connected clients in the workspace",
    "Query the real-time service (via Redis) to determine current table visibility.",
    "",
    "3. Implement Airtable webhook registration (where available):",
    "For Airtable bases, register webhook listeners via Airtable webhook API. When a webhook",
    "fires, enqueue immediate P0 inbound sync job, bypassing polling. Fall back to polling if",
    "webhook registration fails. Store webhook config in base_connections.sync_config.webhooks:",
    "{ airtable_webhook_id, airtable_webhook_cursor, webhook_registered_at }",
    "",
    "4. Add converted table skip logic to the sync dispatcher:",
    "Before dispatching, check base_connection.sync_status:",
    "- 'converted' or 'converted_finalized' -> skip entirely",
    "- 'converted_dual_write' -> dispatch but write to shadow records only",
    "- All other active statuses -> dispatch normally",
    "",
    "5. Update BullMQ repeatable job configuration:",
    "Replace fixed-interval repeatable job with dynamic scheduling:",
    "- Scheduler runs every 30 seconds",
    "- Each run evaluates all active base_connections",
    "- Determine polling interval per table based on visibility",
    "- Enqueue sync job only if enough time elapsed since last poll",
    "- Track last poll time in Redis: sync:last_poll:{baseConnectionId}:{tableId}",
    "",
    "ACCEPTANCE CRITERIA:",
    "- getPollingInterval() returns correct interval for each visibility state",
    "- Scheduler queries Socket.io room membership (via Redis) for table visibility",
    "- Converted tables skipped entirely — no sync jobs dispatched",
    "- converted_dual_write tables dispatch to shadow only",
    "- Scheduler only enqueues when enough time elapsed since last poll",
    "- Airtable webhook registration attempted; webhook syncs bypass polling",
    "- Webhook failure falls back gracefully to polling",
    "- Per-table poll tracking in Redis updated on each sync dispatch",
    "- Unit tests cover all visibility state combinations",
    "- Integration test: mock Socket.io room state, verify correct intervals",
    "- ESLint and TypeScript compile with zero errors",
    "- Coverage >= 80% on new files",
    "",
    "DO NOT BUILD:",
    "- Priority-based scheduling P0-P3 (Prompt 4)",
    "- Multi-tenant fairness (Prompt 4)",
    "- Sync status indicators or dashboard UI (Prompts 5, 10)",
    "- Error recovery flows (Prompts 6-9)",
]
add_colored_container(doc, "PASTE INTO CLAUDE CODE", prompt3_lines, BLUE_BG)

add_body(doc, "Wait for Claude Code to finish. You should see the new scheduler file and passing tests.")

add_colored_container(doc, "GIT COMMAND", [
    "git add apps/worker/src/jobs/sync-scheduler.ts \\",
    "       apps/worker/src/jobs/sync-inbound.ts \\",
    "       packages/shared/sync/types.ts",
    "git add -A",
    'git commit -m "feat(sync): implement smart polling with adaptive intervals and converted table skip logic [Phase 2C, Prompt 3]"',
], GREEN_BG)

add_separator(doc)
doc.add_page_break()

# ============================================================
# INTEGRATION CHECKPOINT 1
# ============================================================
add_heading(doc, "INTEGRATION CHECKPOINT 1 (after Prompts 1–3)", level=1)

add_colored_container(doc, "CONTEXT — READ FOR UNDERSTANDING", [
    "A stop-and-verify point. You're checking that everything from the first 3 prompts works",
    "together correctly before moving on. Think of it like a quality inspection on a construction",
    "site — you don't put up the walls until you've verified the foundation is solid.",
], GRAY_BG, mono=False)

add_colored_container(doc, "VERIFICATION — PASTE INTO CLAUDE CODE", [
    "Run these verification commands and report the results:",
    "1. pnpm turbo typecheck — must show zero errors",
    "2. pnpm turbo lint — must show zero errors",
    "3. pnpm turbo test — all tests must pass",
    "4. pnpm turbo test -- --coverage — coverage thresholds must be met",
    "5. Verify that Notion field type registrations are present in FieldTypeRegistry",
    "   alongside Airtable registrations — no conflicts between the two adapters",
    "6. Verify that the smart polling scheduler correctly transitions between",
    "   intervals when table visibility changes",
], ORANGE_BG)

add_body(doc, "If anything fails: Tell Claude Code what failed. It will attempt to fix the issue. Do not move on until all 6 checks pass.")

add_colored_container(doc, "GIT COMMAND", [
    "git add -A",
    'git commit -m "chore(verify): integration checkpoint 1 — Notion adapter + smart polling [Phase 2C, CP-1]"',
    "git push -u origin feat/phase-2c-notion-error-recovery-dashboard",
], GREEN_BG)

add_body(doc, "Review Point: Go to GitHub and look at your branch. You should see all the new Notion adapter files and the smart scheduler. CI checks should be running (or passed). If CI fails, come back to Claude Code and fix the issues before continuing.")

add_separator(doc)
doc.add_page_break()

# ============================================================
# PROMPT 4
# ============================================================
add_heading(doc, "PROMPT 4: Add Priority-Based Sync Scheduling", level=1)

add_context_block(doc,
    "When you edit a record in EveryStack and need it synced back to Airtable/Notion, that's urgent — it should happen immediately. But a background check for new data on a table nobody's looking at? That can wait. This prompt adds a priority system with 4 levels. Critical user actions (P0) always go through. Background polling (P3) only runs when there's plenty of bandwidth. It also ensures fairness — one busy customer can't hog all the sync capacity.",
    "A new priority module will be created. Tests will verify that P0 jobs always dispatch while P3 jobs are skipped when the system is busy. The multi-tenant fairness system will ensure round-robin scheduling.",
    "5–8 minutes."
)

prompt4_lines = [
    "Build Prompt 4: Priority-Based Job Scheduling (P0-P3) and Multi-Tenant Fairness.",
    "",
    "Depends on: Prompt 3 (smart polling scheduler).",
    "",
    "Load context from: sync-engine.md lines 504-531 (Priority-Based Scheduling,",
    "Multi-Tenant Fairness, Sync Status UI).",
    "",
    "Target files:",
    "- apps/worker/src/jobs/sync-scheduler.ts (extend)",
    "- apps/worker/src/jobs/sync-priority.ts (new)",
    "- packages/shared/sync/types.ts (extend)",
    "",
    "No migration required.",
    "",
    "TASK:",
    "Layer priority-based scheduling on top of the adaptive polling scheduler and enforce",
    "multi-tenant fairness.",
    "",
    "1. Define the priority tier system:",
    "enum SyncPriority {",
    "  P0_CRITICAL = 0,  // Outbound sync (cell edits), webhook-triggered inbound",
    "  P1_ACTIVE = 1,    // Inbound polling for actively viewed tables",
    "  P2_BACKGROUND = 2, // Inbound polling for non-visible tables",
    "  P3_INACTIVE = 3,  // Inbound polling for inactive workspaces",
    "}",
    "",
    "2. Build apps/worker/src/jobs/sync-priority.ts:",
    "export function evaluatePriority(",
    "  priority: SyncPriority, capacityPercent: number",
    "): PriorityDecision",
    "Rules:",
    "  P0: Always dispatch.",
    "  P1: Dispatch if >30% capacity remains. Delayed if <30%.",
    "  P2: Dispatch if >50% capacity remains. Skipped if <50%.",
    "  P3: Dispatch if >70% capacity remains. Skipped if <70%.",
    "",
    "3. Query remaining rate limit capacity:",
    "export async function getRateLimitCapacity(",
    "  platform: string, scopeKey: string",
    "): Promise<number>  // Returns 0-100 percentage remaining",
    "Read the ZSET for the rate limit key, count tokens in current window, compare against",
    "the registered limit.",
    "",
    "4. Implement multi-tenant fairness:",
    "Within each priority tier, tenants are served round-robin.",
    "- Per-platform round-robin index in Redis: sync:rr:{platform}:{priorityTier}",
    "- Per-tenant poll budget: single tenant max 20% of platform rate limit capacity",
    "- P0 is exempt from the 20% cap (critical user-initiated operations always proceed)",
    "- Plan tier does NOT affect sync freshness — fairness is absolute across all plans",
    "export async function getNextTenantForPlatform(",
    "  platform, priorityTier, eligibleTenants",
    "): Promise<string | null>",
    "",
    "5. Integrate priority into the scheduler:",
    "Update sync-scheduler.ts to:",
    "- Assign SyncPriority based on source (outbound edit -> P0, active poll -> P1, etc.)",
    "- Before dispatching, call evaluatePriority() with current capacity",
    "- Log P1-P3 delays/skips via Pino for observability",
    "",
    "6. Wire outbound sync jobs as P0:",
    "Update outbound sync job (from 2B) to always use P0_CRITICAL.",
    "Webhook-triggered inbound syncs also use P0.",
    "",
    "ACCEPTANCE CRITERIA:",
    "- SyncPriority enum defines four tiers (P0-P3)",
    "- evaluatePriority() correctly dispatches/delays/skips based on capacity thresholds",
    "- getRateLimitCapacity() correctly reads remaining capacity from Redis ZSET",
    "- Multi-tenant round-robin dispatches jobs fairly within each priority tier",
    "- Per-tenant 20% capacity cap enforced for P1-P3 (not P0)",
    "- Outbound sync jobs always dispatched as P0",
    "- Webhook-triggered inbound syncs dispatched as P0",
    "- P1-P3 delays/skips logged via Pino",
    "- Unit tests cover all priority x capacity combinations",
    "- Integration test: simulate capacity pressure, verify P0 dispatches while P3 skipped",
    "- ESLint and TypeScript compile with zero errors",
    "- Coverage >= 80% on new files",
    "",
    "DO NOT BUILD:",
    "- Sync status UI indicators (Prompt 5)",
    "- Error recovery flows (Prompts 6-9)",
    "- Dashboard or management UI (Prompt 10)",
    "- Auto-scaling of rate limits based on plan tier",
]
add_colored_container(doc, "PASTE INTO CLAUDE CODE", prompt4_lines, BLUE_BG)

add_colored_container(doc, "GIT COMMAND", [
    "git add apps/worker/src/jobs/sync-scheduler.ts \\",
    "       apps/worker/src/jobs/sync-priority.ts \\",
    "       packages/shared/sync/types.ts",
    "git add -A",
    'git commit -m "feat(sync): implement priority-based job scheduling P0-P3 and multi-tenant fairness [Phase 2C, Prompt 4]"',
], GREEN_BG)

add_separator(doc)
doc.add_page_break()

# ============================================================
# PROMPT 5
# ============================================================
add_heading(doc, "PROMPT 5: Build Sync Health Tracking and Status Badges", level=1)

add_context_block(doc,
    "Every sync connection needs a health monitor — like a dashboard light in your car. This prompt creates the health tracking system (is the connection healthy? stale? erroring? needing re-authentication?) and builds the visual badge that appears in the table header showing the current sync status. Users will see things like \"Synced 2 min ago\" (green), \"Sync retrying...\" (yellow), or \"Re-authentication required\" (red).",
    "New type definitions for connection health, a health derivation utility, a data function for fetching sync status, and a React badge component. Tests will cover all 8 possible health states.",
    "5–8 minutes."
)

prompt5_lines = [
    "Build Prompt 5: ConnectionHealth/SyncError Types, Sync Status Indicators UI,",
    "and Staleness Threshold.",
    "",
    "This prompt is independent — operates on existing base_connections.health JSONB column.",
    "",
    "Load context from: sync-engine.md lines 818-870 (Sync Connection Status Model,",
    "Sync Status Indicators), field-groups.md lines 312-326 (Sync Status Indicator icon).",
    "",
    "Target files:",
    "- packages/shared/sync/types.ts (extend — ConnectionHealth, SyncError)",
    "- packages/shared/sync/health.ts (new)",
    "- apps/web/src/components/sync/SyncStatusBadge.tsx (new)",
    "- apps/web/src/data/sync-status.ts (new)",
    "",
    "No migration required.",
    "",
    "TASK:",
    "Define health tracking types, build health status derivation utility, create sync status",
    "badge component.",
    "",
    "1. Define ConnectionHealth and SyncError types in packages/shared/sync/types.ts:",
    "ConnectionHealth: { last_success_at, last_error (SyncError | null),",
    "  consecutive_failures, next_retry_at, records_synced, records_failed }",
    "SyncError: { code (SyncErrorCode), message, timestamp, retryable, details }",
    "SyncErrorCode: 'auth_expired' | 'rate_limited' | 'platform_unavailable' |",
    "  'schema_mismatch' | 'permission_denied' | 'partial_failure' |",
    "  'quota_exceeded' | 'unknown'",
    "Add Zod schemas for ConnectionHealth and SyncError to validate JSONB on read.",
    "",
    "2. Build packages/shared/sync/health.ts:",
    "Health states: 'healthy' | 'syncing' | 'stale' | 'retrying' | 'error' |",
    "  'auth_required' | 'paused' | 'conflicts'",
    "export function deriveSyncHealthState(",
    "  syncStatus, health, lastSyncAt, pollingInterval,",
    "  pendingConflictCount, isSyncing",
    "): SyncHealthState",
    "Staleness threshold = 2 x pollingInterval.",
    "Also build: export function updateConnectionHealth(",
    "  existing, event: 'sync_success' | 'sync_error', details?",
    "): ConnectionHealth",
    "On success: reset consecutive_failures, clear last_error.",
    "On error: increment consecutive_failures, set last_error.",
    "",
    "3. Build apps/web/src/data/sync-status.ts:",
    "export async function getSyncStatusForTable(",
    "  tenantId, tableId",
    "): Promise<SyncStatus>",
    "Returns derived health state, last sync time, pending conflict count, platform info.",
    "Uses getDbForTenant().",
    "",
    "4. Build apps/web/src/components/sync/SyncStatusBadge.tsx:",
    "React component for sync status in table header:",
    "- healthy: 'Synced 2 min ago' — textSecondary (green dot)",
    "- syncing: 'Syncing...' — accent (teal, animated pulsing)",
    "- stale: 'Last synced 45 min ago' — warning (yellow dot)",
    "- retrying: 'Sync retrying...' — warning (yellow, pulsing), tooltip shows next retry",
    "- error: 'Sync error' — error (red), click opens sync settings",
    "- auth_required: 'Re-authentication required' — error (red), click navigates to re-auth",
    "- paused: 'Sync paused' — textSecondary (gray)",
    "- conflicts: 'Synced · N conflicts' — warning (amber dot), click opens conflict resolution",
    "Use shadcn/ui Badge and Tooltip. Follow design system color tokens.",
    "Use lightweight relative time formatter — no heavyweight library.",
    "",
    "ACCEPTANCE CRITERIA:",
    "- ConnectionHealth and SyncError types exported from packages/shared/sync/types.ts",
    "- Zod schemas validate ConnectionHealth and SyncError JSONB shapes",
    "- deriveSyncHealthState() correctly derives all 8 health states",
    "- Staleness threshold = 2 x pollingInterval",
    "- updateConnectionHealth() correctly transitions on success and error",
    "- getSyncStatusForTable() returns correct sync status using getDbForTenant()",
    "- testTenantIsolation() passes for getSyncStatusForTable()",
    "- SyncStatusBadge renders all 8 health states with correct colors and text",
    "- Badge click navigates to sync settings for error states",
    "- Badge click navigates to re-auth for auth_required",
    "- Relative time display shows correct human-readable time",
    "- ESLint and TypeScript compile with zero errors",
    "- Coverage >= 80% on new files",
    "",
    "DO NOT BUILD:",
    "- Error recovery flow implementations (Prompts 6-9)",
    "- Sync settings dashboard (Prompt 10)",
    "- Synced table sidebar badges (Prompt 11)",
    "- Notification system (Prompt 9)",
]
add_colored_container(doc, "PASTE INTO CLAUDE CODE", prompt5_lines, BLUE_BG)

add_colored_container(doc, "GIT COMMAND", [
    "git add packages/shared/sync/types.ts packages/shared/sync/health.ts \\",
    "       apps/web/src/components/sync/SyncStatusBadge.tsx \\",
    "       apps/web/src/data/sync-status.ts",
    "git add -A",
    'git commit -m "feat(sync): implement ConnectionHealth types, sync status indicators, and staleness threshold [Phase 2C, Prompt 5]"',
], GREEN_BG)

add_separator(doc)
doc.add_page_break()

# ============================================================
# PROMPT 6
# ============================================================
add_heading(doc, "PROMPT 6: Build Error Recovery (Auth, Permissions, Outages)", level=1)

add_context_block(doc,
    "When a sync connection breaks, users need to know what happened and how to fix it. This prompt handles the three most common failure scenarios: (1) The user's login token expired — they see a red banner and a \"Re-authenticate\" button. (2) The platform revoked permissions — they're told to fix it on the platform side. (3) The platform is down (Airtable/Notion outage) — the system retries automatically with increasing delays (1 min, 5 min, 15 min, up to 6 hours). No data is lost during any of these.",
    "A centralized error handler that classifies sync errors, a re-auth banner component, a server action for handling re-authentication, and exponential backoff logic. Tests will cover all error classification paths.",
    "8–12 minutes."
)

prompt6_lines = [
    "Build Prompt 6: Auth Expired, Permission Denied, and Platform Unavailable Recovery Flows.",
    "",
    "Depends on: Prompt 5 (ConnectionHealth/SyncError types, updateConnectionHealth()).",
    "",
    "Load context from: sync-engine.md lines 871-942 (Error Recovery Flows: Auth Expired,",
    "Platform Unavailable sections).",
    "",
    "Target files:",
    "- apps/worker/src/jobs/sync-error-handler.ts (new)",
    "- apps/web/src/components/sync/ReauthBanner.tsx (new)",
    "- apps/web/src/actions/sync-reauth.ts (new)",
    "",
    "No migration required.",
    "",
    "TASK:",
    "Implement three error recovery flows for the most common sync failure scenarios.",
    "",
    "1. Build apps/worker/src/jobs/sync-error-handler.ts:",
    "Centralized error handler invoked when a sync attempt fails:",
    "export async function handleSyncError(",
    "  baseConnectionId, error, context: SyncJobContext",
    "): Promise<void>",
    "This function:",
    "- Classifies error into one of the 8 SyncErrorCode categories",
    "- Calls updateConnectionHealth() to update base_connections.health",
    "- Transitions sync_status based on error type",
    "- Schedules retry (retryable) or marks for manual intervention",
    "Error classification:",
    "- HTTP 401/403 -> auth_expired or permission_denied",
    "- HTTP 429 -> rate_limited (auto-recovery via backoff)",
    "- HTTP 5xx or timeout -> platform_unavailable",
    "- Validation errors on individual records -> partial_failure",
    "- Schema structure changes -> schema_mismatch",
    "",
    "2. Auth expired recovery flow:",
    "When sync receives 401/403:",
    "- Set sync_status = 'auth_required' on base_connection",
    "- Set health.last_error = { code: 'auth_expired', retryable: false }",
    "- No further sync attempts until re-authenticated",
    "Build apps/web/src/components/sync/ReauthBanner.tsx:",
    "- Red banner at top of affected table",
    "- Text: 'Your connection to {Platform} has expired. Re-authenticate to resume syncing.'",
    "- Button: [Re-authenticate] -> triggers OAuth re-auth flow",
    "Build apps/web/src/actions/sync-reauth.ts:",
    "- Handles re-auth completion callback",
    "- Stores new OAuth token in base_connections.credentials",
    "- Resets sync_status = 'active', clears health.last_error",
    "- Enqueues immediate P0 catch-up sync from last_sync_at",
    "No data loss: local edits accumulate normally during auth downtime.",
    "",
    "3. Permission denied recovery:",
    "Similar to auth expired, different message:",
    "- health.last_error = { code: 'permission_denied', retryable: false }",
    "- Banner: 'Your {Platform} integration no longer has write access. Ask the {Platform}",
    "  admin to restore permissions.'",
    "- No automatic recovery — user fixes on platform side, then clicks [Retry Now]",
    "",
    "4. Platform unavailable recovery:",
    "When platform returns 5xx or times out:",
    "- Retry with exponential backoff: 1m -> 5m -> 15m -> 1h -> 3h -> 6h",
    "- Track in health.next_retry_at and health.consecutive_failures",
    "- During retries: badge shows 'Sync retrying...' (yellow, pulsing)",
    "- After 6 hours continuous failure: set sync_status = 'error'",
    "- Badge: 'Sync error — {Platform} may be down. [Retry Now] [Pause Sync]'",
    "Backoff schedule: [60_000, 300_000, 900_000, 3_600_000, 10_800_000, 21_600_000]",
    "export function getBackoffDelay(consecutiveFailures): number | null",
    "'Retry Now' resets consecutive_failures to 0, enqueues immediate sync.",
    "'Pause Sync' sets sync_status = 'paused'.",
    "",
    "ACCEPTANCE CRITERIA:",
    "- handleSyncError() correctly classifies errors into 8 SyncErrorCode categories",
    "- Auth expired: 401/403 -> sync_status = 'auth_required', no further sync attempts",
    "- ReauthBanner displays correct platform name and re-auth button",
    "- Re-auth completion: new token stored, sync_status reset, P0 catch-up sync enqueued",
    "- Catch-up sync uses last_sync_at cursor — no data loss",
    "- Permission denied: appropriate message, manual retry via [Retry Now]",
    "- Platform unavailable: exponential backoff with correct schedule",
    "- health.next_retry_at set correctly during backoff",
    "- After 6 hours of failures: sync_status = 'error'",
    "- 'Retry Now' resets backoff and enqueues immediate sync",
    "- 'Pause Sync' sets sync_status = 'paused'",
    "- All health transitions use updateConnectionHealth()",
    "- testTenantIsolation() passes for any new data access functions",
    "- ESLint and TypeScript compile with zero errors",
    "- Coverage >= 80% on new files",
    "",
    "DO NOT BUILD:",
    "- Partial failure recovery (Prompt 7)",
    "- Schema mismatch recovery (Prompt 8)",
    "- Quota exceeded recovery (Prompt 9)",
    "- Notification system (Prompt 9)",
    "- Sync dashboard (Prompt 10)",
]
add_colored_container(doc, "PASTE INTO CLAUDE CODE", prompt6_lines, BLUE_BG)

add_colored_container(doc, "GIT COMMAND", [
    "git add apps/worker/src/jobs/sync-error-handler.ts \\",
    "       apps/web/src/components/sync/ReauthBanner.tsx \\",
    "       apps/web/src/actions/sync-reauth.ts",
    "git add -A",
    'git commit -m "feat(sync): implement auth expired, permission denied, and platform unavailable recovery flows [Phase 2C, Prompt 6]"',
], GREEN_BG)

add_separator(doc)
doc.add_page_break()

# ============================================================
# INTEGRATION CHECKPOINT 2
# ============================================================
add_heading(doc, "INTEGRATION CHECKPOINT 2 (after Prompts 4–6)", level=1)

add_colored_container(doc, "CONTEXT — READ FOR UNDERSTANDING", [
    "Another quality inspection. You're verifying that the priority scheduling, health tracking,",
    "and error recovery all work together correctly.",
], GRAY_BG, mono=False)

add_colored_container(doc, "VERIFICATION — PASTE INTO CLAUDE CODE", [
    "Run these verification commands and report the results:",
    "1. pnpm turbo typecheck — must show zero errors",
    "2. pnpm turbo lint — must show zero errors",
    "3. pnpm turbo test — all tests must pass",
    "4. pnpm turbo test -- --coverage — coverage thresholds must be met",
    "5. Verify that priority scheduling correctly dispatches P0 jobs even when",
    "   capacity is below 30%",
    "6. Verify that SyncStatusBadge renders correctly for all 8 health states",
    "7. Verify that the error handler correctly classifies mock 401, 429, and 5xx errors",
], ORANGE_BG)

add_body(doc, "If anything fails: Tell Claude Code what failed and let it fix the issues. Do not proceed until all 7 checks pass.")

add_colored_container(doc, "GIT COMMAND", [
    "git add -A",
    'git commit -m "chore(verify): integration checkpoint 2 — priority scheduling + error recovery [Phase 2C, CP-2]"',
    "git push origin feat/phase-2c-notion-error-recovery-dashboard",
], GREEN_BG)

add_body(doc, "Review Point: Check GitHub. CI should be green. Review the diff to make sure the error handler and priority system look reasonable.")

add_separator(doc)
doc.add_page_break()

# ============================================================
# PROMPT 7
# ============================================================
add_heading(doc, "PROMPT 7: Handle Individual Record Failures", level=1)

add_context_block(doc,
    "Sometimes when syncing a batch of 100 records, 98 succeed but 2 fail (maybe the data is in an unexpected format, or a field was too large). Instead of failing the entire batch, this prompt teaches the system to save the 98 good records and track the 2 failures separately. Failed records get automatically retried up to 3 times. If they still fail, the user sees them in a \"Failures\" tab where they can retry individually, skip them, or edit the data to fix the problem.",
    "The sync pipeline will handle partial failures gracefully. A new Failures tab component will display failed records with retry/skip buttons. Tests will verify that good records aren't lost when bad records fail.",
    "8–12 minutes."
)

prompt7_lines = [
    "Build Prompt 7: Partial Failure Recovery Flow with sync_failures Population",
    "and Failures Tab UI.",
    "",
    "Depends on: Prompt 5 (ConnectionHealth types), Prompt 6 (error handler framework).",
    "",
    "Load context from: sync-engine.md lines 943-1000 (Partial Failure recovery flow,",
    "sync_failures table schema, auto-retry policy).",
    "",
    "Target files:",
    "- apps/worker/src/jobs/sync-error-handler.ts (extend)",
    "- packages/shared/sync/sync-failures.ts (new)",
    "- apps/web/src/components/sync/FailuresTab.tsx (new)",
    "- apps/web/src/data/sync-failures.ts (new)",
    "- apps/web/src/actions/sync-failure-actions.ts (new)",
    "",
    "No migration required — sync_failures table already exists from Phase 1B schema.",
    "",
    "TASK:",
    "Implement partial failure recovery: when some records in a sync batch fail while",
    "others succeed, commit the successes and track the failures for resolution.",
    "",
    "1. Update inbound sync pipeline for partial failure handling:",
    "In sync-inbound job, wrap each record's toCanonical() and database write in try/catch.",
    "When a record fails:",
    "- Write it to sync_failures table with error details and failing payload",
    "- Continue processing remaining records in the batch",
    "- After batch completes, set health.records_failed and mark 'completed_with_errors'",
    "The sync pipeline must NOT fail the entire batch for individual record failures.",
    "",
    "2. Build packages/shared/sync/sync-failures.ts:",
    "Data access functions:",
    "- createSyncFailure(tenantId, failure)",
    "- getSyncFailuresForConnection(tenantId, baseConnectionId)",
    "- retrySyncFailure(tenantId, failureId)",
    "- skipSyncFailure(tenantId, failureId, resolvedBy)",
    "- bulkRetrySyncFailures(tenantId, baseConnectionId)",
    "- bulkSkipSyncFailures(tenantId, baseConnectionId, resolvedBy)",
    "",
    "3. Implement auto-retry policy:",
    "Partial failures auto-retry up to 3 times across subsequent sync cycles:",
    "- Each sync cycle checks for sync_failures with status = 'pending', retry_count < 3",
    "- Re-attempt using stored payload, increment retry_count",
    "- After 3 failures: set status = 'requires_manual_resolution', stop auto-retrying",
    "",
    "4. Build apps/web/src/data/sync-failures.ts:",
    "export async function getSyncFailures(",
    "  tenantId, baseConnectionId",
    "): Promise<SyncFailureWithRecord[]>",
    "Join sync_failures with records to include record display name. Uses getDbForTenant().",
    "",
    "5. Build apps/web/src/components/sync/FailuresTab.tsx:",
    "Failures tab content for the Sync Settings Dashboard. For each failure:",
    "- Record name (linked to open record view)",
    "- Error description (human-readable, no technical codes)",
    "- Action buttons: [Retry], [Skip], [Edit in EveryStack]",
    "Bulk actions: [Retry All], [Skip All]",
    "Use shadcn/ui Card, Button, Badge, ScrollArea primitives.",
    "",
    "6. Retention policy:",
    "Resolved failures retained 30 days (cleanup via scheduled BullMQ job).",
    "Pending failures retained indefinitely.",
    "",
    "7. Build apps/web/src/actions/sync-failure-actions.ts:",
    "Server Actions: retrySyncFailureAction, skipSyncFailureAction,",
    "bulkRetrySyncFailuresAction, bulkSkipSyncFailuresAction",
    "",
    "ACCEPTANCE CRITERIA:",
    "- Inbound sync handles individual record failures without failing entire batch",
    "- Failed records written to sync_failures with correct error code, message, payload",
    "- Successfully synced records committed normally",
    "- Sync status shows 'completed_with_errors' when some records fail",
    "- Auto-retry processes pending failures (retry_count < 3) on subsequent sync cycles",
    "- After 3 retries, status set to 'requires_manual_resolution'",
    "- getSyncFailures() returns failures with record display names",
    "- testTenantIsolation() passes for getSyncFailures(), retrySyncFailure(), skipSyncFailure()",
    "- FailuresTab renders failure list with error descriptions and action buttons",
    "- [Retry] re-enqueues the specific failure record",
    "- [Skip] marks failure as resolved",
    "- [Retry All] and [Skip All] bulk actions work correctly",
    "- Resolved failures have resolved_at timestamp set",
    "- ESLint and TypeScript compile with zero errors",
    "- Coverage >= 80% on new files",
    "",
    "DO NOT BUILD:",
    "- Schema mismatch recovery (Prompt 8)",
    "- Quota exceeded recovery (Prompt 9)",
    "- Full sync settings dashboard layout (Prompt 10)",
    "- '[Map to Field...]' action for schema mismatches",
]
add_colored_container(doc, "PASTE INTO CLAUDE CODE", prompt7_lines, BLUE_BG)

add_colored_container(doc, "GIT COMMAND", [
    "git add packages/shared/sync/sync-failures.ts \\",
    "       apps/web/src/components/sync/FailuresTab.tsx \\",
    "       apps/web/src/data/sync-failures.ts \\",
    "       apps/web/src/actions/sync-failure-actions.ts",
    "git add -A",
    'git commit -m "feat(sync): implement partial failure recovery with sync_failures population and Failures tab UI [Phase 2C, Prompt 7]"',
], GREEN_BG)

add_separator(doc)
doc.add_page_break()

# ============================================================
# PROMPT 8
# ============================================================
add_heading(doc, "PROMPT 8: Handle Schema Changes from External Platforms", level=1)

add_context_block(doc,
    "Sometimes a user renames a field, changes a field type, or deletes a field on Airtable or Notion — outside of EveryStack. The sync system needs to detect these changes and present them to the workspace manager for resolution. For example: \"Field 'Status' type changed from Text to Select on Airtable. Accept this change?\" This prevents silent data corruption and gives managers full control.",
    "A schema change detector, a Schema Changes tab component with accept/reject actions for each type of change, and tests covering all 4 change types (type changed, field deleted, field added, field renamed).",
    "8–12 minutes."
)

prompt8_lines = [
    "Build Prompt 8: Schema Mismatch Recovery Flow and Schema Sync Change Detection.",
    "",
    "Depends on: Prompt 5 (ConnectionHealth types), Prompt 6 (error handler framework).",
    "",
    "Load context from: sync-engine.md lines 1001-1065 (Schema Mismatch recovery flow,",
    "sync_schema_changes table, Schema Sync section, impact analysis).",
    "",
    "Target files:",
    "- packages/shared/sync/schema-change-detector.ts (new)",
    "- packages/shared/sync/sync-schema-changes.ts (new)",
    "- apps/worker/src/jobs/sync-inbound.ts (extend)",
    "- apps/web/src/components/sync/SchemaChangesTab.tsx (new)",
    "- apps/web/src/data/sync-schema-changes.ts (new)",
    "- apps/web/src/actions/sync-schema-actions.ts (new)",
    "",
    "No migration required — sync_schema_changes table exists from Phase 1B schema.",
    "",
    "TASK:",
    "Implement schema change detection during inbound sync and the Schema Changes",
    "resolution panel.",
    "",
    "1. Build packages/shared/sync/schema-change-detector.ts:",
    "During each inbound sync, compare platform's current field schema against locally",
    "stored synced_field_mappings:",
    "export function detectSchemaChanges(",
    "  localMappings, platformFields",
    "): SchemaChange[]",
    "Detection logic:",
    "- Field type changed: mapped field's platform type differs from stored type",
    "- Field deleted: field in synced_field_mappings has no match on platform",
    "- Field added: platform field has no matching synced_field_mappings row",
    "- Field renamed: mapped field's name differs (match on platform_field_id)",
    "",
    "2. Build packages/shared/sync/sync-schema-changes.ts:",
    "Data access functions:",
    "- createSchemaChange(tenantId, change)",
    "- getSchemaChangesForConnection(tenantId, baseConnectionId)",
    "- acceptSchemaChange(tenantId, changeId, resolvedBy)",
    "- rejectSchemaChange(tenantId, changeId, resolvedBy)",
    "",
    "3. Compute downstream impact analysis:",
    "export async function computeSchemaChangeImpact(",
    "  tenantId, fieldId",
    "): Promise<SchemaChangeImpact>",
    "SchemaChangeImpact: { formulaCount, automationCount, portalFieldCount, crossLinkCount }",
    "Store computed impact in sync_schema_changes.impact JSONB.",
    "",
    "4. Integrate into inbound sync pipeline:",
    "At start of each inbound sync cycle, run detectSchemaChanges(). If changes detected:",
    "- Write to sync_schema_changes with status = 'pending'",
    "- For field_type_changed and field_deleted: pause syncing affected fields until resolved",
    "- For field_added: continue syncing — data arrives but no local field definition yet",
    "- Show banner: 'Schema change detected on {Platform}. [Review Changes]'",
    "",
    "5. Build apps/web/src/components/sync/SchemaChangesTab.tsx:",
    "Schema Changes tab for Sync Settings Dashboard. For each pending change:",
    "- Field type changed: Warning icon. 'Field type changed: {old} -> {new}.",
    "  Affects {N} formulas.' Actions: [Accept Change] | [Reject — Keep Local Type]",
    "- Field deleted: Error icon. 'Field deleted on {Platform}. {N} records with data.",
    "  Data preserved locally.' Actions: [Archive Field] | [Delete Field]",
    "- Field added: Success icon. 'New field added on {Platform}.'",
    "  Actions: [Add to EveryStack] | [Ignore]",
    "- Field renamed: Info icon. 'Field renamed to {new_name}.'",
    "  Actions: [Accept Rename] | [Keep Local Name]",
    "Show impact analysis below type-changed and deleted field entries.",
    "",
    "6. Implement resolution actions via Server Actions:",
    "- Accept Change (type changed): Update fields.field_type and",
    "  synced_field_mappings.platform_field_type. Re-validate existing data. Resume syncing.",
    "- Reject (type changed): Keep local type. Mark field as local-only.",
    "- Archive Field (deleted): Soft-archive — data preserved, field hidden from grid.",
    "- Delete Field (deleted): Delete field definition and clear canonical data.",
    "- Add to EveryStack (new): Create fields row, synced_field_mappings row, start syncing.",
    "- Ignore (new): Mark schema change as rejected.",
    "- Accept Rename: Update fields.name.",
    "- Keep Local Name: Reject the change.",
    "All resolutions set status, resolved_at, and resolved_by.",
    "",
    "7. Deleted fields are soft-archived (not deleted) to prevent broken cross-links.",
    "",
    "ACCEPTANCE CRITERIA:",
    "- detectSchemaChanges() correctly identifies all 4 change types",
    "- Schema changes written to sync_schema_changes with correct data",
    "- Impact analysis computes formula, automation, portal, cross-link counts",
    "- Affected fields paused during sync until Manager resolves",
    "- New fields continue syncing",
    "- SchemaChangesTab renders all 4 change types with correct icons and actions",
    "- 'Accept Change' updates field type and resumes syncing",
    "- 'Archive Field' soft-archives — data preserved, hidden from grid",
    "- 'Add to EveryStack' creates field and mapping, starts syncing",
    "- All resolutions set resolved_at and resolved_by",
    "- testTenantIsolation() passes for all new data access functions",
    "- ESLint and TypeScript compile with zero errors",
    "- Coverage >= 80% on new files",
    "",
    "DO NOT BUILD:",
    "- Partial failure recovery (Prompt 7)",
    "- Quota exceeded recovery (Prompt 9)",
    "- Full sync settings dashboard layout (Prompt 10)",
    "- Formula engine or field type migration tooling",
]
add_colored_container(doc, "PASTE INTO CLAUDE CODE", prompt8_lines, BLUE_BG)

add_colored_container(doc, "GIT COMMAND", [
    "git add packages/shared/sync/schema-change-detector.ts \\",
    "       packages/shared/sync/sync-schema-changes.ts \\",
    "       apps/web/src/components/sync/SchemaChangesTab.tsx \\",
    "       apps/web/src/data/sync-schema-changes.ts \\",
    "       apps/web/src/actions/sync-schema-actions.ts",
    "git add -A",
    'git commit -m "feat(sync): implement schema mismatch recovery and schema sync change detection [Phase 2C, Prompt 8]"',
], GREEN_BG)

add_separator(doc)
doc.add_page_break()

# ============================================================
# PROMPT 9
# ============================================================
add_heading(doc, "PROMPT 9: Handle Quota Limits and Build Sync Notifications", level=1)

add_context_block(doc,
    "Two things. First, when a workspace hits its record limit (e.g., 250,000 records on the Professional plan), the sync system now handles it gracefully — it syncs what it can, tells the user how many records couldn't be synced, and offers 4 clear options (add a filter, upgrade the plan, delete old records, or disable some tables). Second, it builds the notification system for all sync issues — in-app toasts for immediate problems, emails for critical issues like expired auth or sustained outages.",
    "A quota exceeded panel component, a notification pipeline that sends toasts and emails, and an escalation check job that monitors for sustained downtime.",
    "8–12 minutes."
)

prompt9_lines = [
    "Build Prompt 9: Quota Exceeded Recovery Flow and Sync Notification System.",
    "",
    "Depends on: Prompt 5 (ConnectionHealth types).",
    "",
    "Load context from: sync-engine.md lines 1066-1097 (Quota Exceeded recovery flow),",
    "sync-engine.md lines 880-920 (Notification System for Sync Issues).",
    "",
    "Target files:",
    "- apps/worker/src/jobs/sync-error-handler.ts (extend)",
    "- apps/web/src/components/sync/QuotaExceededPanel.tsx (new)",
    "- packages/shared/sync/sync-notifications.ts (new)",
    "- apps/realtime/src/handlers/sync-notifications.ts (new)",
    "",
    "No migration required.",
    "",
    "TASK:",
    "Implement quota exceeded recovery flow and the notification system for all sync issues.",
    "",
    "1. Quota exceeded recovery flow:",
    "When inbound sync detects record quota would be exceeded:",
    "- Sync records that fit (partial sync — up to remaining quota)",
    "- Set sync_status = 'error', health.last_error.code = 'quota_exceeded'",
    "- Record unsynced count in health.last_error.details.unsyncedCount",
    "",
    "2. Build apps/web/src/components/sync/QuotaExceededPanel.tsx:",
    "Panel shown when Manager clicks 'Record quota reached' badge:",
    "- Shows current count vs plan limit (e.g., '250,000 of 250,000 records')",
    "- Shows how many records couldn't be synced",
    "- 4 resolution options:",
    "  [Add Sync Filter] -> navigates to Tables & Filters tab",
    "  [Upgrade Plan] -> navigates to billing settings",
    "  [Delete Records] -> navigates to table with filter for oldest records",
    "  [Disable Tables] -> navigates to Tables & Filters tab with table toggles",
    "- After freeing quota, [Resume Sync] triggers immediate sync cycle",
    "",
    "3. Build the sync notification system:",
    "Notification pipeline for sync issues:",
    "- Conflict detected (manual mode): in-app badge + toast -> Table Manager(s), immediate",
    "- Auth expired: in-app banner + email via Resend -> Owner + Admins, immediate",
    "- 3 consecutive sync failures: in-app toast -> Table Manager(s), after 3rd failure",
    "- Sync down >1 hour: email via Resend -> Workspace Owner, after 1 hour",
    "- Sync down >6 hours: email (escalation) -> Owner + Admins, after 6 hours",
    "- Partial failure (>10 records): in-app toast + badge -> Table Manager(s), after sync",
    "- Schema mismatch detected: in-app banner on table -> Table Manager(s), immediate",
    "- Rate limit sustained (>5 min): in-app subtle indicator, no notifications (auto-resolves)",
    "",
    "4. Build packages/shared/sync/sync-notifications.ts:",
    "export async function sendSyncNotification(",
    "  tenantId, event, details",
    "): Promise<void>",
    "This function:",
    "- Determines recipients based on event type and workspace roles",
    "- For in-app: publish to Redis channel t:{tenantId}:user:{userId}:notifications",
    "- For email: enqueue BullMQ job that sends via Resend",
    "- Use toast-only for now (persistent notifications table is Post-MVP)",
    "",
    "5. Wire notification triggers into error handler:",
    "Update sync-error-handler.ts to call sendSyncNotification() at trigger points:",
    "- On conflict detected (manual resolution mode)",
    "- On auth expired (immediate)",
    "- On 3rd consecutive failure (health.consecutive_failures === 3)",
    "- On sustained downtime (check last_success_at against 1h and 6h thresholds)",
    "- On partial failure with >10 records failed",
    "- On schema mismatch detected",
    "",
    "6. Escalation tier check job:",
    "Scheduled BullMQ job every 15 minutes. Checks all active base_connections:",
    "- If last_success_at >1 hour ago and sync_status !== 'paused': send 1-hour email",
    "- If last_success_at >6 hours ago: send escalation email to Owner + Admins",
    "",
    "ACCEPTANCE CRITERIA:",
    "- Quota exceeded: records that fit are synced, remaining tracked in unsyncedCount",
    "- QuotaExceededPanel renders current quota, plan limit, and 4 resolution options",
    "- [Resume Sync] triggers immediate sync cycle after quota is freed",
    "- Notification system sends in-app toasts for immediate events",
    "- Auth expired triggers email to Owner + Admins via Resend",
    "- 3rd consecutive failure triggers in-app toast to Table Managers",
    "- Escalation check job runs every 15 minutes, sends emails for >1h and >6h downtime",
    "- Notifications published to correct Redis channel for real-time forwarding",
    "- Rate-limited events do NOT generate notifications",
    "- testTenantIsolation() passes for any new data access functions",
    "- ESLint and TypeScript compile with zero errors",
    "- Coverage >= 80% on new files",
    "",
    "DO NOT BUILD:",
    "- Persistent notifications table — Post-MVP. Toast-only for now.",
    "- Notification preferences UI — Post-MVP",
    "- Push notifications (mobile) — Phase 3H",
    "- Full sync settings dashboard layout (Prompt 10)",
]
add_colored_container(doc, "PASTE INTO CLAUDE CODE", prompt9_lines, BLUE_BG)

add_colored_container(doc, "GIT COMMAND", [
    "git add apps/worker/src/jobs/sync-error-handler.ts \\",
    "       apps/web/src/components/sync/QuotaExceededPanel.tsx \\",
    "       packages/shared/sync/sync-notifications.ts \\",
    "       apps/realtime/src/handlers/sync-notifications.ts",
    "git add -A",
    'git commit -m "feat(sync): implement quota exceeded recovery flow and sync notification system [Phase 2C, Prompt 9]"',
], GREEN_BG)

add_separator(doc)
doc.add_page_break()

# ============================================================
# PROMPT 10
# ============================================================
add_heading(doc, "PROMPT 10: Build the Sync Settings Dashboard", level=1)

add_context_block(doc,
    "This is the control center for sync connections — a full management dashboard with 6 tabs. Overview shows the connection's health at a glance. Tables & Filters lets managers turn tables on/off and set up data filters. Conflicts shows any data conflicts waiting to be resolved. Failures shows records that failed to sync. Schema Changes shows field changes detected from the platform. History shows a log of past sync runs. Plus three action buttons: \"Sync Now\" (force an immediate sync), \"Pause Sync\" (stop syncing), and \"Disconnect\" (remove the connection entirely). All the tab contents were built in earlier prompts — this prompt assembles them into the dashboard layout.",
    "A new route and page at /[workspaceId]/settings/sync/[baseConnectionId] with a tabbed interface. All 6 tabs will render with their content. The action buttons will work. Tab badges will show counts of pending issues.",
    "10–15 minutes."
)

prompt10_lines = [
    "Build Prompt 10: Sync Settings Dashboard (6-Tab Layout).",
    "",
    "Depends on: Prompt 5 (sync status types), Prompt 7 (FailuresTab),",
    "Prompt 8 (SchemaChangesTab), Prompt 9 (QuotaExceededPanel).",
    "",
    "Load context from: sync-engine.md lines 1035-1070 (Sync Settings Dashboard wireframe,",
    "Sync History tab, Tables & Filters tab, Sync Now button).",
    "",
    "Target files:",
    "- apps/web/src/app/(workspace)/[workspaceId]/settings/sync/[baseConnectionId]/page.tsx",
    "- apps/web/src/components/sync/SyncDashboard.tsx (new)",
    "- apps/web/src/components/sync/OverviewTab.tsx (new)",
    "- apps/web/src/components/sync/TablesFiltersTab.tsx (new)",
    "- apps/web/src/components/sync/ConflictsTab.tsx (new)",
    "- apps/web/src/components/sync/HistoryTab.tsx (new)",
    "- apps/web/src/data/sync-dashboard.ts (new)",
    "- apps/web/src/actions/sync-dashboard-actions.ts (new)",
    "",
    "No migration required.",
    "",
    "TASK:",
    "Build the Sync Settings Dashboard — central management surface for sync connections.",
    "",
    "1. Create the route and page:",
    "(workspace)/[workspaceId]/settings/sync/[baseConnectionId]/page.tsx",
    "Accessed via: table header sync status badge, workspace settings, sidebar sync icon.",
    "",
    "2. Build apps/web/src/components/sync/SyncDashboard.tsx:",
    "Dashboard layout with 6 tabs using shadcn/ui Tabs.",
    "Header: 'Sync Connection: {Platform} -> {Base Name}'",
    "Summary: Status, Last Sync, Direction (Bidirectional), Polling Interval,",
    "  Conflicts count, Failures count",
    "Tabs: [Overview] [Tables & Filters] [Conflicts (N)] [Failures (N)]",
    "  [Schema Changes (N)] [History]",
    "Tab badges show pending counts (red badge if >0).",
    "Action buttons: [Sync Now] [Pause Sync] [Disconnect]",
    "",
    "3. Build OverviewTab.tsx:",
    "Connection summary and sync history sparkline:",
    "- Connection status, platform, base name, direction",
    "- Last sync time and record count",
    "- 7-day sync history bar chart (green=success, yellow=partial, red=failed)",
    "- Average duration and average records per sync",
    "",
    "4. Build TablesFiltersTab.tsx:",
    "Per-table row: table name, enabled/disabled toggle, synced/total record count,",
    "  filter summary. Expand arrow opens full filter builder (reuse FilterRule[] from 2A).",
    "Changes trigger 'Save & Re-sync': save filter config -> P0 re-sync -> orphan detection.",
    "",
    "5. Build ConflictsTab.tsx:",
    "Pending conflicts for this connection. Reuses conflict resolution modal from 2B.",
    "List: record name, field name, local vs remote value, [Resolve] button.",
    "Bulk actions: [Resolve All: Keep EveryStack] / [Resolve All: Keep Remote]",
    "",
    "6. Wire in existing tabs:",
    "- Failures tab: Import FailuresTab from Prompt 7",
    "- Schema Changes tab: Import SchemaChangesTab from Prompt 8",
    "",
    "7. Build HistoryTab.tsx:",
    "Table of recent sync runs: Timestamp, Direction, Records synced/failed, Duration, Status.",
    "Click row -> expand to show detailed change counts (created/updated/deleted).",
    "",
    "8. Dashboard actions:",
    "- [Sync Now]: Enqueue immediate P0 sync job. Disabled while syncing. Shows 'Syncing...'",
    "- [Pause Sync]: Sets sync_status = 'paused'.",
    "- [Resume Sync]: Resets sync_status = 'active', enqueues catch-up sync.",
    "- [Disconnect]: Confirmation dialog. Removes credentials, stops sync, marks tables as",
    "  local-only. Does NOT delete synced data.",
    "",
    "9. Build apps/web/src/data/sync-dashboard.ts:",
    "- getSyncDashboardData(tenantId, baseConnectionId): Promise<SyncDashboardData>",
    "- getSyncHistory(tenantId, baseConnectionId, days): Promise<SyncHistoryEntry[]>",
    "",
    "ACCEPTANCE CRITERIA:",
    "- Dashboard route accessible at /[workspaceId]/settings/sync/[baseConnectionId]",
    "- 6 tabs render correctly",
    "- Tab badges show pending counts with red indicators",
    "- Overview tab shows connection summary and 7-day sparkline",
    "- Tables & Filters tab shows per-table toggles and filter config with 'Save & Re-sync'",
    "- Conflicts tab lists pending conflicts with resolution actions",
    "- History tab shows recent sync runs with expandable detail rows",
    "- [Sync Now] enqueues P0 sync job, disabled while syncing",
    "- [Pause Sync] sets sync_status = 'paused'",
    "- [Disconnect] removes credentials and marks tables local-only (with confirmation)",
    "- testTenantIsolation() passes for getSyncDashboardData() and getSyncHistory()",
    "- All UI uses shadcn/ui primitives and follows design system tokens",
    "- ESLint and TypeScript compile with zero errors",
    "- Coverage >= 80% on new files",
    "",
    "DO NOT BUILD:",
    "- Real-time dashboard updates (future — data refreshes on page load/manual refresh)",
    "- Export sync history as CSV (Post-MVP)",
    "- Multi-connection comparison view (Post-MVP)",
    "- SmartSuite-specific dashboard elements (Phase 3)",
]
add_colored_container(doc, "PASTE INTO CLAUDE CODE", prompt10_lines, BLUE_BG)

add_colored_container(doc, "GIT COMMAND", [
    'git add "apps/web/src/app/(workspace)/[workspaceId]/settings/sync/" \\',
    "       apps/web/src/components/sync/ \\",
    "       apps/web/src/data/sync-dashboard.ts \\",
    "       apps/web/src/actions/sync-dashboard-actions.ts",
    "git add -A",
    'git commit -m "feat(sync): implement sync settings dashboard with 6-tab layout [Phase 2C, Prompt 10]"',
], GREEN_BG)

add_separator(doc)
doc.add_page_break()

# ============================================================
# PROMPT 11
# ============================================================
add_heading(doc, "PROMPT 11: Add Platform Badges and Sync Icons to the Sidebar", level=1)

add_context_block(doc,
    "When a table is synced from an external platform, users should see that at a glance in the sidebar. This prompt adds two visual indicators: (1) A tiny platform logo (Airtable or Notion) overlaid on the table icon — so you instantly know which platform that table comes from. (2) A sync status icon next to the table name — green arrows if healthy, yellow if there are conflicts, red if something's broken. Native EveryStack tables show neither indicator.",
    "Synced tables in the sidebar will show a small platform logo on their icon and a sync status indicator. Native tables will look exactly the same as before. Clicking the sync status icon will navigate to the Sync Settings Dashboard.",
    "5–8 minutes."
)

prompt11_lines = [
    "Build Prompt 11: Synced Table Tab Badges — Platform Badge and Sync Status Indicator.",
    "",
    "Depends on: Prompt 5 (sync health state derivation).",
    "",
    "Load context from: field-groups.md lines 276-326 (Synced Table Tab Badges —",
    "platform badge, sync status indicator, 6 health states, click actions,",
    "badge + tab color independence).",
    "",
    "Target files:",
    "- apps/web/src/components/sidebar/PlatformBadge.tsx (new)",
    "- apps/web/src/components/sidebar/SyncStatusIcon.tsx (new)",
    "- apps/web/src/components/sidebar/TableTabItem.tsx (extend)",
    "",
    "No migration required.",
    "",
    "TASK:",
    "Add platform badges and sync status indicators to synced table tabs in the sidebar.",
    "",
    "1. Build apps/web/src/components/sidebar/PlatformBadge.tsx:",
    "A 14x14px platform logo overlay at bottom-right of the table type icon:",
    "- Airtable: Airtable logo mark (simplified for 14px)",
    "- Notion: Notion logo mark (black)",
    "- SmartSuite: SmartSuite logo mark (include slot, ships Phase 3)",
    "- null (native): No badge — absence of badge = native EveryStack table",
    "Badge has 1px #FFFFFF border. Position: absolute; bottom: -2px; right: -2px.",
    "Platform logos should be SVG icons. Store in apps/web/src/components/icons/platforms/.",
    "",
    "2. Build apps/web/src/components/sidebar/SyncStatusIcon.tsx:",
    "Sync status indicator to the right of table name in sidebar:",
    "- Healthy: bidirectional arrows, textSecondary, 'Synced with {Platform}. Last sync: {time}'",
    "- Syncing now: arrows (animated rotation), accent (teal), 'Syncing with {Platform}...'",
    "- Conflicts pending: arrows + amber dot, 'N sync conflicts. Click to resolve.'",
    "- Sync paused: paused bars, textSecondary, 'Sync paused. Click to resume.'",
    "- Sync error: arrows + X, error color, 'Sync failed: {reason}. Click for details.'",
    "- Converted: no icon (badge removed after conversion finalized)",
    "Use Lucide React icons. Overlay via CSS pseudo-elements or SVG composites.",
    "Click actions:",
    "- Click sync status icon -> opens Sync Settings Dashboard",
    "- 'Conflicts pending' -> navigates to Dashboard, Conflicts tab",
    "",
    "3. Extend apps/web/src/components/sidebar/TableTabItem.tsx:",
    "If table.base_connection_id is not null -> render PlatformBadge and SyncStatusIcon",
    "If table.base_connection_id is null -> native table, no badge, no sync icon",
    "",
    "4. Badge and tab color independence:",
    "Platform badge and table tab color stripe are independent. A synced Airtable table can",
    "have an amber tab color (user-chosen) AND the Airtable badge. Two pieces of info,",
    "two visual treatments, no conflict. Do not couple badge visibility to tab color.",
    "",
    "ACCEPTANCE CRITERIA:",
    "- PlatformBadge renders correct platform logo at 14x14px with 1px white border",
    "- Native tables render no platform badge",
    "- SyncStatusIcon renders all 6 health states with correct icon, color, tooltip",
    "- 'Syncing now' shows animated rotation on sync icon",
    "- Click on sync status icon navigates to Sync Settings Dashboard",
    "- Click on 'Conflicts pending' navigates to Dashboard -> Conflicts tab",
    "- Converted tables show no sync status icon",
    "- Platform badge and tab color stripe are visually independent",
    "- Platform logos are SVG, crisp at 14px",
    "- Component renders correctly in always-dark sidebar",
    "- ESLint and TypeScript compile with zero errors",
    "- Coverage >= 80% on new files",
    "",
    "DO NOT BUILD:",
    "- Field groups or column groups (Phase 3)",
    "- Board collapse behavior (Phase 3)",
    "- Per-field emphasis or conditional cell coloring (Phase 3)",
    "- SmartSuite platform logo — include component slot but adapter ships Phase 3",
]
add_colored_container(doc, "PASTE INTO CLAUDE CODE", prompt11_lines, BLUE_BG)

add_colored_container(doc, "GIT COMMAND", [
    "git add apps/web/src/components/sidebar/PlatformBadge.tsx \\",
    "       apps/web/src/components/sidebar/SyncStatusIcon.tsx \\",
    "       apps/web/src/components/sidebar/TableTabItem.tsx \\",
    "       apps/web/src/components/icons/platforms/",
    "git add -A",
    'git commit -m "feat(ui): implement synced table tab badges with platform badge and sync status indicator [Phase 2C, Prompt 11]"',
], GREEN_BG)

add_separator(doc)
doc.add_page_break()

# ============================================================
# FINAL INTEGRATION CHECKPOINT
# ============================================================
add_heading(doc, "FINAL INTEGRATION CHECKPOINT (after Prompts 7–11)", level=1)

add_colored_container(doc, "CONTEXT — READ FOR UNDERSTANDING", [
    "The final quality inspection for Phase 2C. This is the most thorough check — you're",
    "verifying that the entire phase works correctly before opening the pull request.",
], GRAY_BG, mono=False)

add_colored_container(doc, "VERIFICATION — PASTE INTO CLAUDE CODE", [
    "Run these verification commands and report the results:",
    "1. pnpm turbo typecheck — zero errors",
    "2. pnpm turbo lint — zero errors",
    "3. pnpm turbo test — all pass",
    "4. pnpm turbo test -- --coverage — thresholds met for all changed packages:",
    "   - packages/shared/sync/ >= 90% lines, 85% branches",
    "   - apps/web/src/data/ >= 95% lines, 90% branches",
    "   - apps/web/src/actions/ >= 90% lines, 85% branches",
    "   - apps/worker/src/jobs/ >= 85% lines, 80% branches",
    "5. pnpm turbo check:i18n — no hardcoded English strings in new UI components",
    "6. Verify Sync Settings Dashboard renders all 6 tabs with correct counts",
    "7. Verify synced table sidebar badges show platform logo and sync status",
    "   for synced tables, no badges for native tables",
    "8. Verify error recovery flows transition health states correctly",
    "   (mock 401 -> auth_required -> re-auth -> active)",
    "9. Verify NotionAdapter field type transforms produce correct canonical",
    "   JSONB shapes for all ~20 Notion property types",
], ORANGE_BG)

add_body(doc, "If anything fails: Tell Claude Code what failed. Fix all issues before proceeding. This is the last gate before the pull request.")

add_colored_container(doc, "GIT COMMAND", [
    "git add -A",
    'git commit -m "chore(verify): final integration checkpoint — Phase 2C complete [Phase 2C, CP-3]"',
    "git push origin feat/phase-2c-notion-error-recovery-dashboard",
], GREEN_BG)

add_separator(doc)
doc.add_page_break()

# ============================================================
# PHASE COMPLETE
# ============================================================
add_heading(doc, "PHASE COMPLETE — Open the Pull Request", level=1)

add_colored_container(doc, "CONTEXT — READ FOR UNDERSTANDING", [
    "You're done building! Now you open a pull request on GitHub so the code can be reviewed",
    "and merged to main. This is like submitting your finished work for approval.",
], GRAY_BG, mono=False)

add_colored_container(doc, "PASTE INTO CLAUDE CODE", [
    "Open a pull request to main with title",
    "'Phase 2C — Notion Adapter, Error Recovery, Sync Dashboard'",
    "and a description summarizing all 11 prompts and 3 checkpoints completed.",
    "",
    "Include the key deliverables:",
    "- NotionAdapter (toCanonical + fromCanonical for ~20 field types)",
    "- Smart polling with adaptive intervals",
    "- Priority-based job scheduling (P0-P3)",
    "- ConnectionHealth tracking with 8 error categories",
    "- Auth/permission/outage/partial-failure/schema-mismatch/quota recovery flows",
    "- Sync notification system with email escalation",
    "- 6-tab Sync Settings Dashboard",
    "- Synced table sidebar badges",
], BLUE_BG)

add_body(doc, "Wait for Claude Code to confirm the PR is created. It will give you a URL.")

add_separator(doc)

add_heading(doc, "After the PR is Merged", level=2)

add_colored_container(doc, "GIT COMMAND", [
    "git checkout main",
    "git pull origin main",
    "git tag v0.2.2-phase-2c",
    "git push origin v0.2.2-phase-2c",
    "git branch -d feat/phase-2c-notion-error-recovery-dashboard",
], GREEN_BG)

add_separator(doc)

add_heading(doc, "What's Next", level=2)

add_colored_container(doc, "CONTEXT — READ FOR UNDERSTANDING", [
    "Phase 3 (MVP — Core UX) is up next. The full sync engine is now operational with two",
    "platform adapters (Airtable + Notion). Phase 3 builds the grid view, card view, record",
    "view, and all user-facing workspace features that consume synced data. The SmartSuite",
    "adapter also ships as part of Phase 3.",
    "",
    "Before starting Phase 3, return to the playbook generation workflow (Session A) to",
    "produce the Phase 3 playbook.",
], GRAY_BG, mono=False)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.expanduser("~/Documents/EveryStack/docs/Playbooks/Phase 2C — Operator Guide.docx")
doc.save(output_path)
print(f"Operator guide saved to: {output_path}")
